# -*- coding: utf-8 -*-
"""
DIR DIVING — MAIN branch complete readiness audit → Word (.docx).

AUDIT ONLY: does not modify application code.
Run: python Docs/generate_main_branch_readiness_audit_full_docx.py

Requires: python-docx
"""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
OUT = HERE / "MAIN_BRANCH_READINESS_AUDIT_FULL_20260519.docx"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"

def _git(*args: str) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=REPO, text=True).strip()
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        return "n/d"


def para(doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_feature_table(doc: Document, title: str, rows: list[tuple[str, ...]]) -> None:
    doc.add_heading(title, level=2)
    ncols = len(rows[0])
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(rows[0]):
        hdr[i].text = h
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    doc = Document()
    head_sha = _git("rev-parse", "HEAD")
    origin_main_sha = _git("rev-parse", "origin/main")
    branch = _git("rev-parse", "--abbrev-ref", "HEAD")

    t = doc.add_heading("DIR DIVING — MAIN BRANCH COMPLETE READINESS AUDIT", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(
        doc,
        "Tipo: audit read-only (nessuna modifica al codice, nessun merge, nessun refactor). "
        "Ambiente host: Windows — Xcode / xcodegen / xcodebuild non disponibili in PATH: "
        "la verifica di compilazione è statica + inferenza da struttura progetto, non build reale.",
        bold=True,
    )
    para(
        doc,
        f"Branch checkout: {branch}. HEAD (locale): {head_sha}. origin/main remoto: {origin_main_sha}. "
        "Nota: il working tree può contenere modifiche non committate; il verdetto compile si basa su `git show HEAD:` per iOS/ContentView.",
    )

    # --- Visual benchmarks ---
    doc.add_heading("Benchmark visivi obbligatori (allegato)", level=1)
    para(
        doc,
        "Watch: stile premium nero/neon DIR DIVING. iOS: companion dark marine con accento cyan. "
        "Le immagini sotto sono i riferimenti forniti per il confronto qualitativo.",
    )
    for cap, path in (
        ("Watch — immersione live (reference)", IMG_WATCH),
        ("iOS — Companion (reference)", IMG_IOS),
    ):
        doc.add_heading(cap, level=2)
        if path.is_file():
            doc.add_picture(str(path), width=Inches(5.7))
        else:
            para(doc, f"[File mancante: {path}]", bold=True)

    doc.add_page_break()

    # Checklist 1–14 (audit input)
    doc.add_heading("Audit checklist (voci richieste 1–14) — esito statico", level=1)
    doc.add_heading("1. Branch and build verification", level=2)
    bullets(
        doc,
        [
            "Branch: main (verificato).",
            "Dipendenze da rami experimental: nessun import diretto nei sorgenti Watch MAIN verso file esclusi da project.yml; "
            "iOS ContentView su HEAD referenzia tipi definiti solo in file esclusi dal target → dipendenza effettiva da file non compilati.",
            "project.yml: sintassi leggibile; validazione YAML automatica non eseguita (modulo PyYAML assente).",
            "XcodeGen / xcodebuild: non eseguiti (tool assenti su Windows).",
            "Bundle ID Watch com.egopfe.dirdiving; iOS com.egopfe.dirdiving.ios; WKCompanionAppBundleIdentifier in Watch Info.plist allineato.",
            "Entitlements: Watch include water-submersion + iCloud+KVS; iOS iCloud senza water-submersion (coerente).",
            "Asset: AppIcon Contents.json presenti per Watch e iOS (verifica statica percorsi).",
            "TODO bloccanti: nessun TODO in file MAIN-only che impedisca build Watch; iOS bloccato strutturalmente prima dei TODO.",
        ],
    )
    doc.add_heading("2. Apple Watch MAIN feature audit", level=2)
    para(doc, "Vedi tabella sezione C (Watch). Copertura funzionale ampia; validazione sensori/GPS solo su device.", bold=False)
    doc.add_heading("3. iOS Companion MAIN feature audit", level=2)
    para(doc, "Vedi tabella sezione C (iOS). Implementazione sorgente presente ma raggiungibilità bloccata da errore di target.", bold=False)
    doc.add_heading("4. UI consistency audit", level=2)
    para(doc, "Dettaglio in sezione E.", bold=False)
    doc.add_heading("5. UX completeness audit", level=2)
    bullets(
        doc,
        [
            "Watch: flussi principali etichettati; pre-dive testo guida GPS; badge aptica disattivata.",
            "iOS: non valutabile end-to-end senza build; rischio tab 'Lab' fuori policy MAIN.",
        ],
    )
    doc.add_heading("6. Settings audit", level=2)
    para(doc, "Dettaglio in sezione F.", bold=False)
    doc.add_heading("7. Haptics and tones audit", level=2)
    para(doc, "Dettaglio in sezione G.", bold=False)
    doc.add_heading("8. Apple Watch hardware interaction audit", level=2)
    para(doc, "Dettaglio in sezione H.", bold=False)
    doc.add_heading("9. Sync audit", level=2)
    para(doc, "Dettaglio in sezione I.", bold=False)
    doc.add_heading("10. Export audit", level=2)
    para(doc, "Dettaglio in sezione J.", bold=False)
    doc.add_heading("11. Safety / disclaimer audit", level=2)
    para(doc, "Dettaglio in sezione K.", bold=False)
    doc.add_heading("12. Error and empty state audit", level=2)
    para(doc, "Dettaglio in sezione L.", bold=False)
    doc.add_heading("13. Code quality / self-consistency audit", level=2)
    para(doc, "Dettaglio in sezione M.", bold=False)
    doc.add_heading("14. Report format", level=2)
    para(doc, "Strutturato in sezioni A–O di questo documento.", bold=False)

    doc.add_page_break()

    # A
    doc.add_heading("A. Branch confermato", level=1)
    bullets(
        doc,
        [
            "Branch corrente (git rev-parse): main.",
            "Target XcodeGen (project.yml): `DIRDiving Watch App` (watchOS 10.0), `DIRDiving iOS` (iOS 17.0).",
            "Build: non eseguita su questo host (assenza Xcode). Esito inferito: vedi sezione A (build status).",
        ],
    )

    # B Executive summary
    doc.add_heading("B. Executive summary", level=1)
    para(
        doc,
        "Percentuali indicative (solo lettura codice + project.yml + commit HEAD; nessuna esecuzione su device).",
        bold=True,
    )
    bullets(
        doc,
        [
            "Readiness complessiva (utente medio su Watch+iPhone): ~48%",
            "Apple Watch MAIN (funzioni core + navigazione): ~82%",
            "iOS Companion MAIN (funzioni dichiarate ma target non coerente): ~35%",
            "UX readiness (coerenza flussi Watch; iOS bloccato a monte): ~62%",
            "Safety / disclaimer readiness: ~74%",
            "Compile readiness (monorepo entrambi i target su HEAD): ~45% — bloccante su iOS",
        ],
    )
    para(
        doc,
        "Blocco principale (CRITICAL): `iOSApp/Views/ContentView.swift` su HEAD importa `ExplorationCenterView` e "
        "`BuddyExperimentalView`, ma `project.yml` esclude quei file dal target iOS. Risultato atteso: errore di "
        "compilazione (simboli non nel modulo).",
        bold=True,
    )

    # C Feature inventory
    doc.add_heading("C. Feature inventory", level=1)
    add_feature_table(
        doc,
        "Watch MAIN — inventario (stato da codice)",
        [
            ("Feature", "Impl", "Reach", "Usable", "Complete", "Notes", "Sev"),
            ("Schermata live immersione", "Y", "Y", "Y", "Y", "DiveLiveView; tema nero/neon", "LOW"),
            ("Profondità corrente", "Y", "Y", "Y", "Y", "Hero numerico", "LOW"),
            ("RunTime", "Y", "Y", "Y", "Y", "Pannello con TTV", "LOW"),
            ("TTV (testo: non safety/NDL)", "Y", "Y", "Y", "Partial", "Copy in Settings; formula non auditata", "MED"),
            ("Cronometro START/STOP/RESET", "Y", "Y", "Y", "Y", "DiveManager + haptic", "LOW"),
            ("Prof. media / max", "Y", "Y", "Y", "Y", "Pannelli secondari", "LOW"),
            ("Temperatura", "Y", "Y", "Y", "Y", "Top bar", "LOW"),
            ("Gauge velocità risalita", "Y", "Y", "Y", "Y", "AscentGaugeView", "LOW"),
            ("Avviso risalita", "Y", "Y", "Y", "Y", "AscentWarningView quando over limit", "LOW"),
            ("Bussola / bearing", "Y", "Y", "Y", "Y", "CompassView tab", "LOW"),
            ("Dive log lista", "Y", "Y", "Y", "Y", "DiveLogListView + empty state", "LOW"),
            ("Dettaglio immersione Watch", "Y", "Y", "Y", "Y", "Export CSV + delete", "LOW"),
            ("GPS inizio / fine", "Y", "Y", "Y", "Y", "Schermate conferma + fallback copy", "MED"),
            ("Export Subsurface CSV", "Y", "Y", "Y", "Y", "ShareLink / temp file Watch", "LOW"),
            ("Immagini utente", "Y", "Y", "Y", "Partial", "UserImagesView — dipende da contenuti", "LOW"),
            ("Impostazioni", "Y", "Y", "Y", "Y", "Scroll + link navigazione", "LOW"),
            ("Allarmi (soglie)", "Y", "Y", "Y", "Y", "AlarmSettingsView; non sync iPhone (documentato)", "MED"),
            ("Info / batteria", "Y", "Y", "Y", "Y", "InfoView da Settings", "LOW"),
            ("Unità (metric)", "Y", "Y", "Y", "Partial", "AppStorage watch units — verificare copertura ovunque", "LOW"),
            ("Aptica / toni", "Y", "Y", "Y", "Y", "HapticService + toggle; WKInterfaceDevice.play", "LOW"),
            ("Digital Crown dedicata", "N", "-", "-", "N", "TabView pagine verticali = Crown implicita; nessuna mappatura custom", "MED"),
            ("Sync verso iPhone", "Y", "Y", "Y", "Y", "WatchSyncService + ack / coda pending", "MED"),
            ("Modalità pre-dive (selector)", "Y", "Y", "Y", "Y", "Solo Diving stabile; copy su experimental isolati", "LOW"),
        ],
    )

    add_feature_table(
        doc,
        "iOS Companion MAIN — inventario (stato da codice + target)",
        [
            ("Feature", "Impl", "Reach", "Usable", "Complete", "Notes", "Sev"),
            ("Logbook", "Y", "Blocked", "N", "N", "Codice presente ma app non compilabile su HEAD", "CRITICAL"),
            ("Dive detail / grafici", "Y", "Blocked", "N", "N", "DiveDetailView.swift presente", "CRITICAL"),
            ("Planner + risultato piano", "Y", "Blocked", "N", "N", "PlannerView con chart Bühlmann", "CRITICAL"),
            ("Gas configuration UI", "Y", "Blocked", "N", "N", "Nella PlannerView", "CRITICAL"),
            ("Curva Bühlmann display", "Y", "Blocked", "N", "N", "Chart(store.buhlmann.curve)", "CRITICAL"),
            ("Export CSV Subsurface", "Y", "Blocked", "N", "N", "SubsurfaceExportService Result-based", "CRITICAL"),
            ("Sync Watch", "Y", "Blocked", "N", "N", "WatchSyncService + MoreView", "CRITICAL"),
            ("Analysis", "Y", "Blocked", "N", "N", "Tab Analysis su ContentView", "CRITICAL"),
            ("Equipment", "Y", "Blocked", "N", "N", "Tab Equipment", "CRITICAL"),
            ("Settings / More", "Y", "Blocked", "N", "N", "DIRWarningBox disclaimer presente", "CRITICAL"),
            ("Cloud backup (iCloud)", "Y", "Blocked", "N", "N", "CloudSyncStore + entitlements", "CRITICAL"),
            ("Explore Lab / Buddy Lab tabs", "Partial", "N", "N", "N", "UI tab su ContentView ma file esclusi dal target", "CRITICAL"),
            ("Notifiche push utente", "N", "N", "N", "N", "Nessun UNUserNotificationCenter trovato", "LOW"),
            ("Toni/haptic iOS", "N", "N", "N", "N", "Nessun UIImpactFeedback in iOSApp", "LOW"),
        ],
    )

    # D Navigation
    doc.add_heading("D. Navigation map", level=1)
    para(doc, "Watch (TabView verticale): ModeSelection → Live → Compass → Settings → UserImages → DiveLog.", bold=False)
    bullets(
        doc,
        [
            "Percorsi Settings: Velocità risalita, Allarmi, Info, Export (testuale), toggle aptica, unità, sync status, retry/clear coda.",
            "Dead end: nessuno evidente; ExportView richiede dismiss esplicito.",
            "Crown: paging tra tab tramite comportamento di sistema; nessuna schermata 'Crown-first' custom.",
        ],
    )
    para(doc, "iOS (HEAD): Tab bar con voci incluse Explore Lab / Buddy Lab — compile mismatch con target.", bold=True)

    # E UI consistency
    doc.add_heading("E. UI consistency vs benchmark", level=1)
    bullets(
        doc,
        [
            "Watch: DiveUI + DiveScreenBackground allineati al reference nero/neon, pannelli arrotondati e bordi sottili — buon match qualitativo.",
            "Watch: rischio clipping solo verificabile su hardware Ultra / display piccoli (non misurato qui).",
            "iOS: DIRTheme (background/surface/cyan) coerente con reference dark marine; card e hairline presenti.",
            "iOS: LogbookView contiene stringa fissa 'MAGGIO 2024' come heading — mismatch UX rispetto a dati reali (severity MED, fix UI-only).",
            "iOS: icone + / … in header logbook senza azioni collegate visibili — possibile dead affordance (MED, verifica wiring).",
        ],
    )

    # F Settings
    doc.add_heading("F. Settings audit", level=1)
    bullets(
        doc,
        [
            "Watch: unità, aptica, allarmi multipli, retry sync, clear coda, stato GPS/sensore/sync — persistenza @AppStorage / UserDefaults.",
            "Watch: riga 'Sync impostazioni' indica locale / bidirezionale planned — aspettativa utente vs realtà (MED, solo chiarezza copy).",
            "iOS: MoreView — sync watch, cloud, demo logbook toggle, export info; coerente ma irraggiungibile se build rotta.",
            "iOS: mancano controlli paralleli Watch per allarmi/aptica (previsto: logica su Watch).",
        ],
    )

    # G Haptics
    doc.add_heading("G. Haptics / tones", level=1)
    bullets(
        doc,
        [
            "Watch: WKInterfaceDevice.play (.success/.notification/.failure/.directionUp/.retry) centralizzati in HapticService; rispettano toggle.",
            "Watch: eventi dive/stopwatch/export/warning collegati in DiveManager, DiveDetailView, DiveLogListView, Settings.",
            "iOS: nessun feedback aptico dedicato trovato — gap UX non safety-critical.",
        ],
    )

    # H Hardware
    doc.add_heading("H. Apple Watch hardware interactions", level=1)
    bullets(
        doc,
        [
            "Side button / Action Button: nessun binding esplicito trovato; fare affidamento su UI tap.",
            "Long press: non mappato globalmente (verificare singole view se necessario).",
            "Gesture conflicts: non rilevati staticamente; richiede prova in immersione.",
        ],
    )

    # I Sync
    doc.add_heading("I. Sync audit (Watch ↔ iPhone)", level=1)
    bullets(
        doc,
        [
            "Watch: WatchSyncService.shared — coda pending, ack, failed count, retry, chiave WatchSyncAuth.",
            "iOS: WatchSyncService riceve messaggi / userInfo / replyHandler (versione su working tree vs HEAD: verificare al merge).",
            "Rischio: divergenza implementazioni iOS/Watch tra commit e working tree — va allineata in fase fix (post-audit).",
        ],
    )

    # J Export
    doc.add_heading("J. Export audit", level=1)
    bullets(
        doc,
        [
            "Watch: CSV Subsurface via writeCSV → URL? + ShareLink in log list / detail.",
            "iOS: writeCSV → Result<URL,ExportError> — gestione errori più esplicita; validità file non testata qui.",
            "GPX/KML: non trovati servizi dedicati in grep rapido — segnalare come assente se richiesto da roadmap.",
        ],
    )

    # K Safety
    doc.add_heading("K. Safety / disclaimer / App Store risk", level=1)
    bullets(
        doc,
        [
            "MoreView iOS: DIRWarningBox — testo 'supporto informativo' (positivo).",
            "Watch Settings: TTV descritto come derivato e non safety/NDL (positivo).",
            "Rischio residuo: assenza build iOS impedisce review completa stringhe in tutte le schermate.",
            "Nessuna evidenza in questo audit che l'app si dichiari 'computer subacqueo certificato' — tono da mantenere.",
        ],
    )

    # L Error / empty
    doc.add_heading("L. Error & empty states", level=1)
    bullets(
        doc,
        [
            "Watch DiveLogListView: empty state + messaggi export.",
            "Watch DiveLiveView: banner errore sensore + GPS pre-dive copy.",
            "iOS: non verificabile runtime; pattern SwiftUI presenti in AnalysisView per sessions.isEmpty su versione estesa (working tree).",
        ],
    )

    # M Code quality
    doc.add_heading("M. Code quality / self-consistency", level=1)
    bullets(
        doc,
        [
            "TODO espliciti concentrati in file experimental (Exploration*, Experimental*) — OK se esclusi dal target; "
            "tab iOS che li referenzia rompe la build: incoerenza architetturale/release.",
            "Due implementazioni SubsurfaceExportService (Watch vs iOS) con API diverse — debito tecnico medio (non bloccante Watch).",
            "BuddyAssist/Haptic paths presenti in codice escluso dal target Watch MAIN — non influiscono se non importati.",
        ],
    )

    # N Bugs table narrative + roadmap
    doc.add_heading("N. Bugs / issue tracker (priorità)", level=1)
    add_feature_table(
        doc,
        "Elenco sintetico",
        [
            ("Title", "Plat", "Severity", "Impact", "Est. fix"),
            ("iOS target esclude file referenziati da ContentView", "iOS", "CRITICAL", "App non compila", "small functional / wiring"),
            ("Tab Explore/Buddy su main vs policy MAIN-only", "iOS", "HIGH", "UX confusione + build", "UI-only + project.yml"),
            ("Logbook heading statico MAGGIO 2024", "iOS", "MED", "Fuorviante", "UI-only"),
            ("Digital Crown non documentata / non custom", "Watch", "LOW", "Aspettative utente Ultra", "docs + optional UX"),
            ("iOS assenza haptic feedback", "iOS", "LOW", "Meno feedback", "UI-only"),
        ],
    )

    doc.add_heading("N2. Priority roadmap", level=1)
    bullets(
        doc,
        [
            "1) Must fix before compile/use: allineare `ContentView` iOS con file inclusi nel target oppure includere voci tab coerenti con `project.yml`.",
            "2) Must fix before TestFlight: build verde entrambi target + smoke test device + Ultra layout spot-check.",
            "3) Must fix before App Store: copy safety completo, privacy labels, export validato, sync stress test.",
            "4) Post-release: Crown mapping avanzato, notifiche iOS, formati export aggiuntivi.",
        ],
    )

    # O Verdict
    doc.add_heading("O. Final verdict", level=1)
    bullets(
        doc,
        [
            "Pronta a compilare (monorepo Watch+iOS su HEAD): NO — iOS rotto per mismatch tab/target.",
            "Pronta per test interno end-to-end: NO (stesso motivo).",
            "Pronta per utente medio (Watch+iPhone): NO finché iOS non compila.",
            "Pronta per TestFlight: NO.",
            "Pronta per App Store: NO.",
            "Cosa blocca il 100%: (1) compile iOS, (2) build reale + QA device, (3) allineamento terminologico e Ultra, (4) export/hardening legale.",
        ],
    )

    doc.add_heading("Appendice — XcodeGen", level=1)
    para(
        doc,
        "project.yml: struttura coerente (bundle IDs, WKCompanionAppBundleIdentifier, embed Watch in iOS, esclusioni experimental). "
        "Validazione YAML automatica non eseguita (PyYAML assente). Consigliato: `xcodegen generate` su macOS + `xcodebuild` per entrambi gli scheme.",
    )

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
