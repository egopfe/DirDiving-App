# -*- coding: utf-8 -*-
"""Genera il roadmap Word (100%) UX/UI Watch + iOS + sicurezza/copy. Esegui: python Docs/generate_ux_roadmap_100_docx.py"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
OUT = HERE / "DIR_DIVING_Piano_100_UX_UI_Watch_iOS_Sicurezza.docx"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"


def p(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return para


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def main():
    doc = Document()
    t = doc.add_heading("DIR DIVING — Piano operativo al 100%", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p(
        doc,
        "Documento: roadmap dettagliata per Watch MAIN, iOS su main, UX Watch e sicurezza/copy "
        "(TTV/TTR, non certificazione, GPS in superficie). "
        "Riferimento visivo: screenshot allegati (immersione live Apple Watch; mockup iOS Companion dark/cyan).",
    )
    p(
        doc,
        "Vincolo assoluto: non modificare funzioni di base né la logica di business (algoritmi decompressione, "
        "modelli gas, calcolo TTV/TTR, campionamento sensori, regole di sync). "
        "Interventi ammessi: validazione build, test su dispositivi (incluso Ultra), copy e disclaimer, "
        "accessibilità, layout SwiftUI (padding, tipografia, contrasto, raggruppamenti), documentazione README/Docs, "
        "coerenza terminologica tra schermate, stati vuoti e messaggi di errore più chiari.",
        bold=True,
    )

    # --- Reference screenshots ---
    doc.add_heading("Allegato visivo — linee guida UI/UX", level=1)
    p(
        doc,
        "Le immagini seguenti definiscono il target di chiarezza e densità informativa. "
        "Ogni intervento proposto in questo documento deve avvicinare il prodotto a questa leggibilità "
        "senza introdurre nuove funzioni o cambiare i calcoli sottostanti.",
    )
    for label, path in (
        ("Watch — schermata immersione live (alta priorità informativa, blocchi delimitati, codice colore funzionale)", IMG_WATCH),
        ("iOS — Companion (dark mode, accento cyan, card, tab, Planner e output piano)", IMG_IOS),
    ):
        doc.add_heading(label, level=2)
        if path.is_file():
            doc.add_picture(str(path), width=Inches(5.8))
            p(doc, f"File: {path.name}", size=9)
        else:
            p(doc, f"[Immagine non trovata: {path} — rigenerare dopo aver copiato ReferenceUI]", bold=True)

    doc.add_page_break()

    # --- Design system derived from references ---
    doc.add_heading("Sintesi del linguaggio visivo di riferimento", level=1)
    bullets(
        doc,
        [
            "Watch: sfondo nero ad alto contrasto; stati positivi in verde (es. “IN IMMERSIONE”, TTV, START); "
            "profondità in bianco dominante con unità in cyan/blu; temperatura con iconografia coerente (goccia).",
            "Watch: metriche critiche in contenitori arrotondati con bordo sottile (verde per blocco TTV/RunTime; giallo per cronometro) "
            "per separare compiti mentali diversi: “quanto manca alla violazione/superficie” vs “tempo di immersione/cronometro”.",
            "Watch: gauge risalita verticale con scala cromatica verde→giallo→rosso; il puntatore deve restare leggibile in condizioni di movimento "
            "(solo rifiniture di spessore/tick/etichette, non soglie algoritmiche).",
            "Watch: area comandi (START / STOP / RESET) con codice colore semantico (verde/rosso/neutro) e icone universali.",
            "iOS: gerarchia a card, titoli di sezione chiari, barra tab cinque voci (Logbook, Analisi, Planner, Attrezzatura, Altro).",
            "iOS: accento cyan su azioni primarie (es. “Calcola piano”) e highlight arancio per consumi/dati “da notare” senza allarmismo inappropriato.",
            "iOS: navigazione a schede secondarie nei dettagli (Riepilogo / Grafici / Dettagli; Piano / Curva / Grafici) per ridurre scroll cognitivo.",
        ],
    )

    doc.add_heading("Area 1 — Watch MAIN (da ~82% al 100%)", level=1)
    p(
        doc,
        "Obiettivo: raggiungere confidenza ingegneristica e di prodotto ripetibile su build watchOS e su Apple Watch Ultra, "
        "allineando la percezione utente allo screenshot di immersione live senza alterare motori di calcolo.",
    )
    doc.add_heading("1.1 Validazione build (ripetibilità)", level=2)
    numbered(
        doc,
        [
            "Documentare in README o Docs uno schema unico: nome workspace, scheme `DIRDiving Watch App`, SDK minimo, passi `xcodegen generate` + `xcodebuild` con destinazione `generic/platform=watchOS`.",
            "Aggiungere una checklist firmabile (data, commit hash, esito) per ogni release candidata.",
            "Se esiste CI: allineare scheme/destinazione a quanto usato in locale; se non esiste CI, valutare solo job di compilazione (nessun deploy) per ridurre regressioni silenziose.",
            "Registrare esito su simulatore watchOS e su hardware reale; annotare differenze di performance solo come note (nessun tuning algoritmico).",
        ],
    )
    doc.add_heading("1.2 Apple Watch Ultra (layout, comfort, lettura)", level=2)
    numbered(
        doc,
        [
            "Verificare che il blocco profondità centrale e il gauge risalita non siano compressi su angoli del display Ultra; "
            "eventuali micro-aggiustamenti solo via SwiftUI (safeAreaInsets, padding, frame min/max).",
            "Validare tipografia a luce solare forte: dove il testo secondario risulta troppo tenue, aumentare contrasto percepito "
            "(colore del testo o peso font), non cambiare i valori numerici mostrati.",
            "Sessione Ultra con polso in movimento: controllare che TTV/RunTime e “IN IMMERSIONE” restino distinguibili a colpo d’occhio "
            "(ordine visivo e spaziatura, non nuove metriche).",
            "Digital Crown / scroll: se presenti liste in immersione o in post-dive, assicurare affordance visiva (hint di scroll) senza nuove interazioni funzionali.",
        ],
    )
    doc.add_heading("1.3 Coerenza con il reference live", level=2)
    numbered(
        doc,
        [
            "Allineare etichette e unità alla stessa densità semantica del reference: titoli brevi sopra, valore dominante, unità disaccoppiata ma vicina.",
            "Evitare stringhe ambigue su stati di immersione: verbi coerenti (“In immersione”, “In pausa”, “Fine immersione”) solo a livello copy.",
            "Gauge risalita: etichette scala leggibili; se un valore è tagliato su piccoli Watch, usare abbreviazioni o rotazione etichetta, non cambiare la mappatura colore-soglia.",
        ],
    )
    doc.add_heading("1.4 Definizione di “100%” per quest’area", level=2)
    bullets(
        doc,
        [
            "Build watchOS verde su destinazione documentata + almeno un dispositivo fisico + Ultra smoke test completato.",
            "Elenco “note Ultra” chiuso o ridotto a differenze accettate documentate.",
        ],
    )

    doc.add_heading("Area 2 — iOS su `main` (completamento post compiler blocker)", level=1)
    p(
        doc,
        "Obiettivo: dopo lo sblocco del compiler blocker, portare iOS al livello di solidità e chiarezza del mockup Companion: "
        "stessa gerarchia a card, stessa chiarezza delle tab e dei flussi Planner → output, senza toccare motori Bühlmann o tabelle deco.",
    )
    doc.add_heading("2.1 Engineering readiness", level=2)
    numbered(
        doc,
        [
            "Ripetere la stessa disciplina di Area 1 per il target `DIRDiving iOS` (xcodegen + xcodebuild documentati).",
            "Verificare coerenza `project.yml` ↔ sorgenti: nessun file referenziato ma escluso dal target.",
            "Smoke test: avvio a freddo, cambio tab, apertura dettaglio immersione, Planner fino a “Calcola piano” (solo navigazione), "
            "More/impostazioni sync Watch se presenti.",
        ],
    )
    doc.add_heading("2.2 Allineamento UI al mockup (solo presentazione)", level=2)
    numbered(
        doc,
        [
            "Logbook: card con immagine sito, metadati (data, ora, max, durata, gas) in griglia leggibile; badge “Buddy” solo se già supportato dal modello — "
            "in caso contrario limitarsi a copy/visibilità del dato esistente.",
            "Dettaglio immersione: tab secondarie per separare riepilogo, grafici e dettagli testuali; evitare muri di testo.",
            "Planner: segmented control (Semplice / Avanzato / Tecnico) con etichette coerenti con il resto dell’app; stepper e card gas come contenitori visivi, non nuove regole MOD/PPO2.",
            "Output piano: tab “Piano / Curva Bühlmann / Grafici” con titoli espliciti; tabelle con intestazioni fisse in scroll orizzontale se necessario (solo layout).",
            "Azioni primarie full-width cyan come nel reference; azioni secondarie (Modifica, Condividi) in stile ghost/outline.",
        ],
    )
    doc.add_heading("2.3 Accessibilità e internazionalizzazione", level=2)
    numbered(
        doc,
        [
            "Dynamic Type: verificare che titoli e valori non si sovrappongano; usare `minimumScaleFactor` o andare a capo, non ridurre contenuto informativo.",
            "VoiceOver: etichette e hint per grafici (profilo, Bühlmann) che descrivono “cosa si sta vedendo”, non interpretazioni mediche.",
            "Allineare tono e termini con Watch (vedi Area 4 e sezione glossario).",
        ],
    )

    doc.add_heading("Area 3 — UX Watch (da ~78% al 100%)", level=1)
    p(
        doc,
        "Obiettivo: massimizzare chiarezza, predictability e feedback percettivo sul polso, replicando la qualità del reference live: "
        "blocchi delimitati, colori semantici, comandi espliciti.",
    )
    doc.add_heading("3.1 Gerarchia e riduzione del carico cognitivo", level=2)
    numbered(
        doc,
        [
            "Per ogni schermata Watch, definire una “frase d’intento” utente (es. “So quanto posso ancora stare e a che profondità sono”).",
            "Raggruppare sempre metriche affine in contenitori (come TTV + RunTime nel reference) per evitare salti visivi casuali.",
            "Usare dimensione font e peso per priorità: profondità attuale > TTV/RunTime > metriche secondarie > testi di stato.",
        ],
    )
    doc.add_heading("3.2 Stati di sistema e messaggi", level=2)
    numbered(
        doc,
        [
            "Loading / errore / successo: messaggi brevi, stesso lessico di iOS quando la funzione è equivalente (sync, export).",
            "Evitare messaggi tecnici (codici errore) senza contesto; aggiungere riga “Cosa fare ora” (es. riavvia sync) senza cambiare la logica di retry.",
        ],
    )
    doc.add_heading("3.3 Controlli e sicurezza d’uso", level=2)
    numbered(
        doc,
        [
            "Pulsanti START/STOP/RESET: conferme leggere solo tramite copy o debounce visivo se già previsto dal design system — nessuna nuova macchina a stati.",
            "Haptics: uniformare intensità per azioni equivalenti (tap conferma vs errore) come rifinitura percettiva.",
        ],
    )

    doc.add_heading("Area 4 — Sicurezza, trasparenza e copy (da ~76% al 100%)", level=1)
    p(
        doc,
        "Obiettivo: massima chiarezza su cosa l’app è e cosa non è; allineare README, stringhe Watch/iOS e eventuali note in Planner "
        "senza alterare calcoli TTV/TTR, SAC, CNS, OTU, ecc.",
    )
    doc.add_heading("4.1 TTV / TTR e altre metriche “sensibili”", level=2)
    numbered(
        doc,
        [
            "Creare un glossario utente (TTV/TTR, SAC, MOD, PPO2, CNS, OTU) con definizione operativa e limite di interpretazione.",
            "Se nel prodotto compaiono entrambe le etichette TTV e TTR, documentare se sono sinonimi di presentazione o metriche distinte; "
            "uniformare le etichette in UI/README solo dopo verifica con il team — senza cambiare il significato numerico già calcolato.",
            "Aggiungere micro-copy contestuale sotto i valori critici (una riga) che rimanda al glossario o a “Scopri di più” statico.",
        ],
    )
    doc.add_heading("4.2 Non certificazione e limiti d’uso", level=2)
    numbered(
        doc,
        [
            "Blocco README + schermata Info: linguaggio chiaro (“strumento di supporto”, “non sostituisce formazione, procedure del dive center, "
            "né dispositivi dedicati omologati”).",
            "Evitare claim assoluti (“garantito”, “certificato”) se non supportati da evidenze; preferire formulazioni condizionali.",
        ],
    )
    doc.add_heading("4.3 GPS in superficie e tracciamento", level=2)
    numbered(
        doc,
        [
            "Spiegare in modo semplice quando il GPS è atteso in superficie, limiti in acqua, accuratezza variabile, impatto batteria — solo testo e link a impostazioni privacy già esistenti.",
            "In Watch live e in iOS dettaglio, allineare le icone e le didascalie (ingresso/uscita) al comportamento reale già implementato.",
        ],
    )
    doc.add_heading("4.4 Planner iOS (solo quando buildabile)", level=2)
    numbered(
        doc,
        [
            "Accanto a output Bühlmann/CNS/OTU: disclaimer breve che richiama i limiti del modello e l’obbligo di doppio controllo umano.",
            "Riordinare sezioni solo se migliora la scansione visiva; non introdurre nuovi passi di calcolo.",
        ],
    )

    doc.add_heading("Coerenza cross-platform e QA lessicale", level=1)
    bullets(
        doc,
        [
            "Tabella di mapping termini Watch ↔ iOS ↔ README (una pagina) da usare in revisione copy.",
            "Sessione di review condivisa: stesso ordine di priorità visiva (profondità, tempo, gas, sicurezza risalita).",
            "Checklist finale: ogni schermata che mostra TTV/TTR, GPS o sync ha almeno una riga di contesto o link a glossario.",
        ],
    )

    doc.add_heading("Milestone suggerite (senza toccare la business logic)", level=1)
    numbered(
        doc,
        [
            "M0 — Compiler blocker iOS risolto + build documentata.",
            "M1 — Watch build + Ultra smoke + note chiuse.",
            "M2 — iOS smoke + allineamento card/tab al reference.",
            "M3 — Copy sicurezza + glossario + README.",
            "M4 — UX Watch pass finale (gerarchia, stati, haptics) + QA lessicale.",
        ],
    )

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
