# -*- coding: utf-8 -*-
"""Generate MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260524_CURRENT_PRE_MODIFICATION.docx."""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
OUT = HERE / "MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260524_CURRENT_PRE_MODIFICATION.docx"
MD = HERE / "MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260524_CURRENT_PRE_MODIFICATION.md"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"


def _git(*args: str) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=REPO, text=True).strip()
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        return "n/d"


def para(doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    doc = Document()
    head = _git("rev-parse", "HEAD")
    branch = _git("rev-parse", "--abbrev-ref", "HEAD")

    title = doc.add_heading(
        "DIR DIVING — UX / Interaction / Feature Accessibility Audit (PRE-MODIFICATION)",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, f"Date: 2026-05-24 · Branch: {branch} · HEAD: {head}", bold=True)
    para(
        doc,
        "Read-only audit. Scope: Watch MAIN + iOS MAIN. No code changes in this pass.",
    )

    doc.add_heading("9. Final Summary", level=1)
    table(
        doc,
        ["Dimension", "Estimate"],
        [
            ["UX / feature accessibility", "~82%"],
            ["Navigation", "~88%"],
            ["Settings", "~78%"],
            ["Hardware", "~75%"],
            ["Haptics", "~85%"],
            ["Sync UX", "~80%"],
            ["Compile readiness", "0% — BUILD FAILED"],
        ],
    )
    para(
        doc,
        "CRITICAL: AscentRateSettingsView.swift:41 — limitControl missing return. Fix before QA.",
        bold=True,
    )

    doc.add_heading("1. Feature Inventory (highlights)", level=1)
    table(
        doc,
        ["Platform", "Feature", "Reachable", "Gap"],
        [
            ["Watch", "Live dive + ascent", "Y", "Device depth"],
            ["Watch", "BUSSOLA", "Y", "—"],
            ["Watch", "ASC SET limits", "P", "Build broken"],
            ["Watch", "7 App Shortcuts", "P", "Not device-tested"],
            ["Watch", "Side button dive", "N", "Shortcuts only"],
            ["iOS", "Planner + ack", "Y", "Persisted"],
            ["iOS", "Logbook/detail/export", "Y", "Units in list OK"],
            ["iOS", "Sync + conflicts", "Y", "Aggregate status only"],
            ["iOS", "Alarms settings", "N", "Watch-local"],
        ],
    )

    doc.add_heading("2. Navigation Map", level=1)
    bullets(
        doc,
        [
            "Watch: Legal → TabView (Live, BUSSOLA, Settings, [Images], Log) → sub-screens.",
            "During dive: only Live, BUSSOLA, Log; Settings blocked.",
            "iOS: Planner | Logbook | Analysis | Equipment | More.",
            "No dead ends found statically.",
        ],
    )

    doc.add_heading("3. Settings Report", level=1)
    bullets(
        doc,
        [
            "Watch: units, language, haptics, ascent, alarms, legal, sync status — OK.",
            "Watch: brightness/tones/export — informational only.",
            "iOS: units, language, sync, cloud, demo, legal — OK.",
            "iOS: no alarm/haptic editors (by design).",
        ],
    )

    doc.add_heading("4. Hardware & Haptics", level=1)
    bullets(
        doc,
        [
            "Crown: vertical paging + steppers in alarm/ascent settings.",
            "Side button: not mapped; WatchShortcutHelpView documents Shortcuts.",
            "7 App Intents registered.",
            "Haptics gated; ascent/depth/confirm/export covered.",
            "No audio tones.",
        ],
    )

    doc.add_heading("5. UX Blockers", level=1)
    table(
        doc,
        ["ID", "Severity", "Issue"],
        [
            ["UX-CR-01", "CRITICAL", "Watch build fail AscentRateSettingsView"],
            ["UX-H-01", "HIGH", "Ultra depth / entitlement QA"],
            ["UX-H-02", "MED", "Side button not mapped"],
            ["UX-M-03", "MED", "Planner metric-only vs units"],
        ],
    )

    doc.add_heading("6–8. Safety, Priority, Code Impact", level=1)
    bullets(
        doc,
        [
            "Safety: disclaimers strong; build blocks device haptic QA.",
            "Immediate: fix compile, then builds + simulator smoke.",
            "Pre-release: Ultra QA, intents, sync on device.",
            "Impact: one small functional fix; rest UI/i18n/QA.",
        ],
    )

    doc.add_heading("Reference UI", level=1)
    if IMG_WATCH.exists():
        doc.add_picture(str(IMG_WATCH), width=Inches(2.0))
    if IMG_IOS.exists():
        doc.add_picture(str(IMG_IOS), width=Inches(2.0))

    para(doc, f"Full markdown: {MD.name}", size=9)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
