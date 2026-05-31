# -*- coding: utf-8 -*-
"""Generate MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_CURRENT.docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
MD = HERE / "MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_CURRENT.md"
OUT = HERE / "MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_CURRENT.docx"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"


def _clean_table_row(line: str) -> list[str]:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return cells


def _is_table_separator(line: str) -> bool:
    stripped = line.strip().replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
    return stripped == ""


def _add_paragraph(doc: Document, text: str, *, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Menlo"
        run.font.size = Pt(9)


def _add_table(doc: Document, rows: list[str]) -> None:
    parsed = [_clean_table_row(row) for row in rows if row.strip()]
    if len(parsed) < 2:
        for row in rows:
            _add_paragraph(doc, row)
        return

    headers = parsed[0]
    body = [row for row in parsed[2:] if len(row) == len(headers)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in body:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_doc() -> Document:
    doc = Document()

    title = doc.add_heading("DIR DIVING — MAIN UX / Interaction / Feature Accessibility Audit", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_paragraph(
        doc,
        "Current pre-modification audit export. Source: MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_CURRENT.md",
        bold=True,
        size=10,
    )

    if IMG_WATCH.exists():
        doc.add_picture(str(IMG_WATCH), width=Inches(2.35))
        caption = doc.add_paragraph("Watch reference UI")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if IMG_IOS.exists():
        doc.add_picture(str(IMG_IOS), width=Inches(2.55))
        caption = doc.add_paragraph("iOS reference UI")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    return doc


def main() -> None:
    if not MD.exists():
        raise FileNotFoundError(f"Missing markdown source: {MD}")

    doc = build_doc()
    lines = MD.read_text(encoding="utf-8").splitlines()

    in_code = False
    code_lines: list[str] = []
    pending_table: list[str] = []

    def flush_table() -> None:
        nonlocal pending_table
        if pending_table:
            _add_table(doc, pending_table)
            pending_table = []

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        if line.startswith("```"):
            flush_table()
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            pending_table.append(line)
            continue

        if pending_table and not line.startswith("|"):
            flush_table()

        if not line.strip():
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.style = "Intense Quote"
            p.add_run(line[2:].strip())
        else:
            _add_paragraph(doc, line, size=10)

    if in_code and code_lines:
        _add_code_block(doc, code_lines)
    flush_table()

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
