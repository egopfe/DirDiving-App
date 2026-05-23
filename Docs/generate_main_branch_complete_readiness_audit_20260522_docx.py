# -*- coding: utf-8 -*-
"""Generate MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260522.docx."""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
OUT = HERE / "MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260522.docx"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"


def _git(*args: str) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=REPO, text=True).strip()
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        return "n/d"


def para(doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    doc = Document()
    head = _git("rev-parse", "HEAD")
    branch = _git("rev-parse", "--abbrev-ref", "HEAD")

    h = doc.add_heading("DIR DIVING — MAIN BRANCH COMPLETE READINESS AUDIT", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, f"Date: 2026-05-22 · Branch: {branch} · HEAD: {head}", bold=True)
    para(
        doc,
        "Audit-only: no code changes. Scope: Apple Watch MAIN + iOS Companion MAIN on unified main. "
        "Re-audit at 800bfa8 supersedes 20260520 report at 9def114.",
    )

    doc.add_heading("A. Branch Confirmed", level=1)
    bullets(
        doc,
        [
            f"Branch: {branch} · HEAD: {head}",
            "project.yml valid; xcodegen generate: PASS.",
            "Experimental sources excluded (Apnea, Snorkeling, Buddy, Exploration).",
            "xcodebuild Watch (Ultra 3 sim) + iOS (iPhone 17): BUILD SUCCEEDED (2026-05-22).",
            "Generic iOS Simulator build: FAILED (duplicate DIR DIVING.app product).",
            "Bundle IDs: com.egopfe.dirdiving / com.egopfe.dirdiving.ios.",
            "Reference UI: Docs/ReferenceUI/*.png present.",
        ],
    )

    doc.add_heading("B. Executive Summary", level=1)
    table(
        doc,
        ["Dimension", "Readiness %"],
        [
            ["Overall MAIN", "83%"],
            ["Apple Watch MAIN", "87%"],
            ["iOS Companion MAIN", "85%"],
            ["UX completeness", "79%"],
            ["Safety / disclaimers", "82%"],
            ["Compile readiness", "90%"],
        ],
    )
    para(
        doc,
        "Verdict: ready for internal/simulator testing and developer-led device QA. "
        "Not 100% for average consumer or App Store until Ultra depth validation, "
        "physical sync QA, i18n on hardcoded screens, and release legal/assets.",
    )

    doc.add_heading("Visual benchmarks", level=2)
    if IMG_WATCH.exists():
        doc.add_picture(str(IMG_WATCH), width=Inches(2.2))
        para(doc, "Watch reference: Watch_LIVE_reference.png")
    if IMG_IOS.exists():
        doc.add_picture(str(IMG_IOS), width=Inches(2.8))
        para(doc, "iOS reference: iOS_Companion_reference.png")

    doc.add_heading("C. Feature Inventory (summary)", level=1)
    table(
        doc,
        ["Platform", "Feature", "Complete", "Severity", "Notes"],
        [
            ["Watch", "Live dive + ascent banner", "Y", "—", "Inline banner; gauge visible"],
            ["Watch", "BUSSOLA / bearing", "Y", "—", "BUSSOLA terminology"],
            ["Watch", "GPS start/end", "Y", "—", "Compact banner ~1.4s"],
            ["Watch", "iPhone inbound WC", "Partial", "MED", "Handlers OK; iOS never pushes sessions"],
            ["Watch", "Settings i18n", "Partial", "MED", "Hardcoded IT blocks"],
            ["iOS", "Logbook / detail / export", "Y", "—", ""],
            ["iOS", "Planner + Bühlmann", "Partial", "MED", "Indicative; share icon inert"],
            ["iOS", "Sync conflicts", "Partial", "MED", "Service only; no UI"],
            ["Both", "Tombstones", "Partial", "MED", "Code OK; device QA needed"],
        ],
    )

    doc.add_heading("D. Navigation Map", level=1)
    para(doc, "Watch: Mode Selection → Live; tabs: Mode | Live | BUSSOLA | Settings | Images | Log.", bold=True)
    para(doc, "iOS: Logbook | Analisi | Planner | Attrezzatura | Altro.", bold=True)
    para(doc, "Unreachable in MAIN: Exploration, Buddy, Experimental views (project.yml excludes).")

    doc.add_heading("E. UI Consistency", level=1)
    bullets(
        doc,
        [
            "Watch: high match to black/neon DiveUI; Settings IT hardcoding lowers score.",
            "iOS: high match to DIRTheme; More disclaimer and Planner IT hardcoded.",
            "Watch AppIcon: unassigned icon_92_2x.png warning.",
        ],
    )

    doc.add_heading("F. Settings", level=1)
    bullets(
        doc,
        [
            "Watch: ascent, alarms, haptics, language — persisted; not synced to iPhone.",
            "iOS: language editable; units/export read-only; cloud sync; demo logbook.",
            "Planner mode persisted but not used in calculations.",
        ],
    )

    doc.add_heading("G. Haptics / Tones", level=1)
    bullets(
        doc,
        [
            "Watch: all paths via HapticService; gated by toggle; ascent/alarm/GPS covered.",
            "iOS: no alert sounds; visual feedback only.",
        ],
    )

    doc.add_heading("H. Hardware", level=1)
    para(
        doc,
        "Crown: vertical TabView paging. Side button not mapped. App Intents: stopwatch promoted; "
        "manual dive/bearing/alarm defined but not in default shortcuts.",
    )

    doc.add_heading("I. Sync", level=1)
    bullets(
        doc,
        [
            "Watch → iPhone: outbound queue + HMAC — implemented.",
            "Tombstones: dirdiving_shared_deleted_session_ids + WC broadcast — implemented.",
            "iPhone → Watch sessions: Watch can ingest; iOS does not send — gap.",
            "iOS sync conflicts: persisted; no resolution UI.",
            "Physical device QA: required before TestFlight.",
        ],
    )

    doc.add_heading("J. Export", level=1)
    para(doc, "Subsurface CSV on Watch and iOS. GPX/KML not on MAIN. More EXPORT card is display-only.")

    doc.add_heading("K. Safety", level=1)
    bullets(
        doc,
        [
            "Not presented as certified dive computer.",
            "Planner/TTV indicative; no mandatory planner ack before calculate.",
            "Water submersion entitlement configured; real Ultra validation required.",
        ],
    )

    doc.add_heading("L. Error / Empty states", level=1)
    para(
        doc,
        "Logbook/Analysis empty states with actions. iOS conflicts silent. Depth may show 0 m without strong no-sensor state."
    )

    doc.add_heading("M. Bugs To Fix (top)", level=1)
    table(
        doc,
        ["#", "Title", "Sev.", "Fix type"],
        [
            ["1", "Depth entitlement not validated on real Ultra", "HIGH", "Process / QA"],
            ["2", "Physical sync + tombstone QA", "HIGH", "Process / QA"],
            ["3", "iPhone does not push sessions to Watch", "MED", "Small functional"],
            ["4", "iOS sync conflict UI missing", "MED", "UI-only"],
            ["5", "Hardcoded IT Settings/More/Planner", "MED", "UI-only"],
            ["6", "Watch AppIcon warning", "MED", "Assets"],
            ["7", "Generic iOS sim build fails", "LOW", "Build config"],
        ],
    )
    para(doc, "Resolved: Watch inbound WC, tombstones, GPS banner, alarm dismiss (a75a6c3).", bold=True)
    para(doc, "Stale: MAIN_BRANCH_ISSUES_AND_PRIORITIES still lists P0-SYNC-01 open — incorrect.", bold=True)

    doc.add_heading("N. Priority Roadmap", level=1)
    bullets(
        doc,
        [
            "Before compile/use: named-simulator builds; portal entitlements.",
            "Before TestFlight: Ultra depth + physical sync/tombstone QA.",
            "Before App Store: legal/privacy, i18n/assets, conflict UI or docs.",
            "Post-release: Mode Selection skip, GPX, settings sync, PlanResult share.",
        ],
    )

    doc.add_heading("O. Final Verdict", level=1)
    table(
        doc,
        ["Question", "Answer"],
        [
            ["Ready to compile?", "Yes (named simulators)"],
            ["Ready for internal test?", "Yes with playbook + device QA gaps"],
            ["Ready for average user?", "Conditional"],
            ["Ready for TestFlight?", "After device depth + sync QA"],
            ["Ready for App Store?", "No"],
        ],
    )

    para(doc, "Full markdown: Docs/MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260522.md", size=9)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
