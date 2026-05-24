# -*- coding: utf-8 -*-
"""Generate MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260524_POST_DEV_NOTES.docx."""
from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
MD = HERE / "MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260524_POST_DEV_NOTES.md"
OUT = HERE / "MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260524_POST_DEV_NOTES.docx"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"


def _git(*args: str) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=REPO, text=True).strip()
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        return "n/d"


def para(doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def add_md_section(doc: Document, lines: list[str], start: int, end: int) -> None:
    """Add markdown lines between headings as paragraphs (simplified)."""
    for line in lines[start:end]:
        s = line.rstrip()
        if not s or s.startswith("```"):
            continue
        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("| ") and "---" not in s:
            continue  # tables handled separately in script
        elif s.startswith("- "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif s.startswith("**") and s.endswith("**"):
            para(doc, s.strip("*"), bold=True)
        else:
            para(doc, re.sub(r"\*\*([^*]+)\*\*", r"\1", s))


def main() -> None:
    head = _git("rev-parse", "--short", "HEAD")
    branch = _git("rev-parse", "--abbrev-ref", "HEAD")
    md_text = MD.read_text(encoding="utf-8") if MD.exists() else ""
    md_lines = md_text.splitlines()

    doc = Document()
    title = doc.add_heading(
        "DIR DIVING — MAIN UX / Interaction / Feature Accessibility Audit",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, f"Post development-notes pass · Date: 2026-05-24 · Branch: {branch} · HEAD: {head}", bold=True)
    para(
        doc,
        "Pre-modification audit only. Scope: Apple Watch MAIN + iOS Companion MAIN. "
        "Experimental targets excluded per project.yml.",
    )

    doc.add_heading("9. Final Summary", level=1)
    table(
        doc,
        ["Dimension", "Estimate", "Notes"],
        [
            ["Release readiness (UX)", "~84%", "Core flows usable on c23d4d4"],
            ["UX completeness", "~80%", "Dev-notes features reachable"],
            ["Stability (interaction)", "~88%", "Sync/conflict mostly surfaced"],
            ["Safety completeness (UX)", "~85%", "Legal/ascent/MOD strong"],
        ],
    )
    para(
        doc,
        "Verdict: Suitable for continued TestFlight. Remaining work is polish (planner units, "
        "launch disclaimer, dive-start affordance), not missing core navigation.",
        bold=True,
    )

    doc.add_heading("1. Feature Inventory (summary)", level=1)
    para(doc, "Watch MAIN — highlights", bold=True)
    bullets(
        doc,
        [
            "Vertical TabView: Live, Compass, Settings, [User Images], Dive Log.",
            "Dive lifecycle sensor-driven; manual fallback when depth API unavailable.",
            "Back affordances on pushed settings, log detail, export, info (c23d4d4).",
            "Units applied on Live and Log via WatchDepthFormatting.",
            "Mode Selection / Apnea / Snorkeling excluded or hidden.",
        ],
    )
    para(doc, "iOS MAIN — highlights", bold=True)
    bullets(
        doc,
        [
            "Tabs: Planner, Logbook, Analysis, Equipment, More.",
            "Equipment templates + editor; planner cylinders/MOD validation.",
            "Manual dive edit reachable for manual sessions.",
            "Planner advanced-only; metric notice shown.",
            "Watch photo preprocess + localized conversion warning.",
        ],
    )

    doc.add_heading("2. Navigation Map", level=1)
    para(
        doc,
        "Watch: NavigationStack → TabView with Settings pushes and log export destinations. "
        "Active dive restricts tabs to Live, Compass, Log. "
        "iOS: TabView per surface; Planner → PlanResult; Logbook → Detail → Manual editor.",
    )

    doc.add_heading("3. Settings Report", level=1)
    table(
        doc,
        ["Platform", "Exposed", "Not in UI / not synced"],
        [
            ["Watch", "Units, language, haptics, alarms, ASC SET, legal", "Depth 35/38/40 bands; alarms not WC-synced"],
            ["iOS", "Units, language, sync, iCloud, demo, legal, pairing reset", "Bulk export (info only); planner ack session"],
        ],
    )

    doc.add_heading("4. Hardware Interaction", level=1)
    bullets(
        doc,
        [
            "Crown: tab paging (system); threshold tuning on Alarms + ASC SET only.",
            "Side button: not mapped; 7 App Intents via Shortcuts (documented).",
            "Haptics: dive, alarms, ascent, depth limits, confirmations — master toggle respected.",
        ],
    )

    doc.add_heading("5. UX Blockers", level=1)
    table(
        doc,
        ["ID", "Severity", "Issue"],
        [
            ["B1", "HIGH", "No Start Dive button when automation active"],
            ["B3", "MEDIUM", "Companion disclaimer every launch"],
            ["B4", "MEDIUM", "Planner metric-only vs global units"],
            ["B5", "MEDIUM", "Disabled planner mode tabs visible"],
            ["B6", "MEDIUM", "Manual editor weak back navigation"],
        ],
    )

    doc.add_heading("6. Safety Issues", level=1)
    bullets(
        doc,
        [
            "No CRITICAL safety UX blockers at c23d4d4.",
            "Max-depth alarm disabled by default — user must enable (MEDIUM).",
            "MOD warnings informative only — Calculate not blocked.",
            "Strong: legal gates, ascent banner, depth limit UI, MOD validation display.",
        ],
    )

    doc.add_heading("7. Recommended Priority", level=1)
    bullets(
        doc,
        [
            "Immediate: dive-start affordance, disclaimer persistence, alarm default UX.",
            "Pre-release: planner units/modes, manual editor toolbar, dive-active log gating.",
            "Post-release: team planner edit, i18n sweep, device QA intents/sync.",
        ],
    )

    doc.add_heading("8. Code Impact", level=1)
    para(
        doc,
        "Most remaining fixes are small UI/product decisions. No architectural rewrite required.",
    )

    doc.add_heading("Reference UI", level=1)
    if IMG_WATCH.exists():
        doc.add_picture(str(IMG_WATCH), width=Inches(2.2))
        para(doc, "Watch reference", size=9)
    if IMG_IOS.exists():
        doc.add_picture(str(IMG_IOS), width=Inches(2.2))
        para(doc, "iOS reference", size=9)

    doc.add_page_break()
    doc.add_heading("Full report (markdown source)", level=1)
    para(doc, f"Complete tables and appendices: {MD.name}", bold=True)
    if md_lines:
        add_md_section(doc, md_lines, 0, min(len(md_lines), 120))

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
