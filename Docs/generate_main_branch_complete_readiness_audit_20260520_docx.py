# -*- coding: utf-8 -*-
"""Generate MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260520.docx from audit markdown."""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
OUT = HERE / "MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260520.docx"
MD = HERE / "MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260520.md"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"


def _git(*args: str) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=REPO, text=True).strip()
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        return "n/d"


def para(doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    doc = Document()
    head = _git("rev-parse", "HEAD")
    branch = _git("rev-parse", "--abbrev-ref", "HEAD")
    dirty = _git("status", "--porcelain")
    dirty_note = " (working tree has uncommitted changes)" if dirty else ""

    title = doc.add_heading("DIR DIVING — MAIN Branch Complete Readiness Audit", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, f"Date: 2026-05-20 · Branch: {branch} · HEAD: {head}{dirty_note}", bold=True)
    para(
        doc,
        "Audit-only (no code changes). Scope: Apple Watch MAIN + iOS Companion MAIN. "
        "Benchmarks: Docs/ReferenceUI/Watch_LIVE_reference.png and iOS_Companion_reference.png.",
    )

    doc.add_heading("A. Branch & Build", level=1)
    bullets(
        doc,
        [
            f"Branch {branch} @ {head}.",
            "xcodegen + Watch + iOS builds: SUCCEEDED (audit run 2026-05-20).",
            "project.yml excludes experimental sources from MAIN targets.",
            "Bundle IDs: iOS com.egopfe.dirdiving.ios; Watch com.egopfe.dirdiving.ios.watch.",
            "Entitlements: iCloud KVS + CloudKit; Watch adds water submersion.",
            "1 non-blocking AppIntents metadata warning on iOS build.",
            "Uncommitted WIP on disk: PlannerSafetyAcknowledgment, CloudSyncStore decode UI, iOS localization.",
        ],
    )

    doc.add_heading("B. Executive Summary", level=1)
    table(
        doc,
        ["Dimension", "db72dce (committed)", "With local WIP"],
        [
            ["Overall", "~91%", "~94%"],
            ["Compile", "100%", "100%"],
            ["Apple Watch MAIN", "~93%", "~93%"],
            ["iOS Companion MAIN", "~90%", "~93%"],
            ["UX completeness", "~93%", "~95%"],
            ["Safety", "~88%", "~88%"],
            ["UI vs reference", "~88%", "~90%"],
        ],
    )

    doc.add_heading("C. Feature Inventory (highlights)", level=1)
    table(
        doc,
        ["Platform", "Feature", "Complete", "Notes"],
        [
            ["Watch", "Live dive + ascent gauge", "Y", "Matches neon reference"],
            ["Watch", "Auto depth / submersion", "P", "Ultra + entitlement required"],
            ["Watch", "7 App Shortcuts", "Y", "Side button via Shortcuts only"],
            ["Watch", "Audio tones", "N", "Informational settings row"],
            ["iOS", "Logbook + detail + CSV", "Y", "No list thumbnails vs mock"],
            ["iOS", "Planner + Bühlmann", "P", "Indicative; not certified DC"],
            ["iOS", "Watch sync + conflicts", "Y", "More UI for conflicts"],
            ["iOS", "iCloud KVS", "P", "Decode error silent until WIP committed"],
            ["iOS", "GPX/KML/.ssrf", "N", "CSV only in MAIN"],
        ],
    )

    doc.add_heading("D. Navigation", level=1)
    bullets(
        doc,
        [
            "Watch: Live → BUSSOLA → Settings → [Images] → Dive Log → Detail.",
            "iOS: Planner → Logbook → Analysis → Equipment → More (reference mock: Logbook first).",
            "No critical dead ends; legal gates intentional.",
        ],
    )

    doc.add_heading("E. UI Consistency vs Reference", level=1)
    table(
        doc,
        ["Platform", "Match", "Gap", "Severity"],
        [
            ["Watch", "Strong", "Some EN labels (RunTime)", "LOW"],
            ["iOS", "Strong theme", "Tab order; no logbook photos", "LOW"],
            ["iOS", "Cards/charts", "Planner/Equipment IT literals", "LOW"],
            ["Marketing", "—", ".ssrf in mock; app ships CSV", "LOW"],
        ],
    )

    doc.add_heading("F–H. Settings / Haptics / Hardware", level=1)
    bullets(
        doc,
        [
            "Watch: units, alarms, ascent, haptics persisted and applied; tones N/A.",
            "iOS: units + demo + cloud + WC sync; planner ack session-only @ db72dce (WIP fixes).",
            "Haptics: comprehensive on Watch; iOS text-only feedback.",
            "Crown: vertical paging; side button not mapped; Action Button via App Intents.",
        ],
    )

    doc.add_heading("I–J. Sync & Export", level=1)
    bullets(
        doc,
        [
            "Watch ↔ iPhone: signed dive payloads, tombstones, units, photos, offline queue.",
            "iCloud: KVS merge; decode failures need WIP UI before App Store.",
            "Export: Subsurface CSV both platforms; ShareLink; GPX/KML not in MAIN.",
        ],
    )

    doc.add_heading("K–L. Safety & Empty States", level=1)
    bullets(
        doc,
        [
            "Strong disclaimers; planner not certified decompression guidance.",
            "Depth 35/38/40 m fixed; ascent banner inline.",
            "Empty states on logbook/analysis; sync/export errors mostly visible.",
            "Blocker: physical Ultra depth QA (R1).",
        ],
    )

    doc.add_heading("M. Bugs To Fix", level=1)
    table(
        doc,
        ["ID", "Severity", "Issue", "Status"],
        [
            ["R1", "HIGH", "Physical depth/submersion QA on Ultra", "Open"],
            ["R2", "MED", "Planner safety ack not persisted", "WIP on disk"],
            ["R3", "MED", "iCloud decode errors silent", "WIP on disk"],
            ["R4", "LOW", "Mixed IT/EN on iOS main screens", "WIP partial"],
            ["R5", "LOW", "Tab order vs reference mock", "Open"],
            ["R6", "LOW", "Logbook thumbnails vs mock", "Open"],
        ],
    )

    doc.add_heading("N. Priority Roadmap", level=1)
    bullets(
        doc,
        [
            "Before compile/use: none (builds succeed).",
            "Before TestFlight: R1 device QA; commit R2–R4 WIP.",
            "Before App Store: R1 + legal/metadata + iCloud error visibility shipped.",
            "Post-release: full i18n, GPX, iOS notifications, UI polish vs reference.",
        ],
    )

    doc.add_heading("O. Final Verdict", level=1)
    table(
        doc,
        ["Question", "db72dce", "After WIP commit"],
        [
            ["Ready to compile?", "YES", "YES"],
            ["Ready for internal test?", "YES", "YES"],
            ["Ready for average user?", "MOSTLY (~91%)", "MOSTLY (~94%)"],
            ["Ready for TestFlight?", "YES after R1", "YES after R1"],
            ["Ready for App Store?", "CONDITIONAL", "CONDITIONAL"],
            ["Blocks 100%", "R1 hardware QA, store review, R2–R4 on remote", "Primarily R1 + metadata"],
        ],
    )

    doc.add_heading("Prior fixes on main", level=1)
    bullets(
        doc,
        [
            "876bcd2: manual edit, merge fields, planner mock row, disclaimer persist, sync fixes.",
            "db72dce: ascent gauge imperial, 7 App Shortcuts, shortcut help, detail refresh.",
        ],
    )

    doc.add_heading("Reference UI", level=1)
    if IMG_WATCH.exists():
        doc.add_picture(str(IMG_WATCH), width=Inches(2.2))
        para(doc, "Watch LIVE reference", size=9)
    if IMG_IOS.exists():
        doc.add_picture(str(IMG_IOS), width=Inches(2.2))
        para(doc, "iOS Companion reference", size=9)

    para(doc, f"Full markdown: {MD.name}", size=9)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
