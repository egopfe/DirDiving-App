# -*- coding: utf-8 -*-
"""Generate MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260524.docx."""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
OUT = HERE / "MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260524.docx"
MD = HERE / "MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260524.md"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"


def _git(*args: str) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=REPO, text=True).strip()
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        return "n/d"


def para(doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    doc = Document()
    head = _git("rev-parse", "HEAD")
    branch = _git("rev-parse", "--abbrev-ref", "HEAD")

    title = doc.add_heading(
        "DIR DIVING — MAIN Branch UX / Interaction / Feature Accessibility Audit",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, f"Date: 2026-05-24 · Branch: {branch} · HEAD: {head}", bold=True)
    para(
        doc,
        "Pre-modification audit only — no code changes. Scope: Apple Watch MAIN + iOS Companion MAIN. "
        "Experimental branches and excluded targets (Snorkeling, Apnea, Buddy) not audited.",
    )

    doc.add_heading("9. Final Summary", level=1)
    table(
        doc,
        ["Dimension", "Estimate", "Notes"],
        [
            ["Release readiness (UX)", "~78%", "Core flows usable; planner/manual gaps"],
            ["UX completeness", "~72%", "Unreachable edit + display inconsistencies"],
            ["Stability (interaction)", "~85%", "Sync/conflict mostly surfaced"],
            ["Safety completeness (UX)", "~80%", "Strong ascent/legal; planner mock row risk"],
        ],
    )
    para(
        doc,
        "Verdict: Continue TestFlight testing on MAIN. Resolve 3 CRITICAL blockers before claiming "
        "planner/manual-dive production-complete.",
        bold=True,
    )

    doc.add_heading("1. Feature Inventory (highlights)", level=1)
    para(doc, "Watch MAIN", bold=True)
    bullets(
        doc,
        [
            "Live, BUSSOLA, Settings, Log, conditional User Images — Crown vertical tabs.",
            "Mode Selection implemented but hidden (single stable mode).",
            "Units picker does not drive Live depth labels.",
            "Runtime alarm UI default 30 min vs engine fallback 60 min until key set.",
        ],
    )
    para(doc, "iOS MAIN", bold=True)
    bullets(
        doc,
        [
            "Tabs: Planner (default), Logbook, Analysis, Equipment, More.",
            "Manual dive add OK; edit implemented but unreachable.",
            "Planner result includes hardcoded ascent row; share icon non-functional.",
            "Launch companion disclaimer on every cold start.",
        ],
    )

    doc.add_heading("2. Navigation Map", level=1)
    para(
        doc,
        "Watch: NavigationStack → TabView (Live, Compass, Settings, [Images], Log) with Settings pushes. "
        "iOS: TabView per surface; Logbook → Detail | Manual add only; Planner → PlanResult.",
    )

    doc.add_heading("3. Settings Report", level=1)
    table(
        doc,
        ["Platform", "Exposed", "Hidden / partial"],
        [
            ["Watch", "Units, language, haptics, ascent limits, alarms, legal", "Skip mode selection; depth 35/38/40; always-on"],
            ["iOS", "Units, language, sync, cloud, demo, legal in More", "resetPairingTrust; planner ack session-only"],
        ],
    )

    doc.add_heading("4. Hardware Interaction", level=1)
    bullets(
        doc,
        [
            "Digital Crown: Watch tab paging + scroll (system).",
            "Side button: not mapped (documented).",
            "Haptics: ascent, depth limits, alarms, stopwatch, compass — respect toggle.",
            "App Intents: only Toggle/Reset stopwatch in Shortcuts catalog (Watch).",
        ],
    )

    doc.add_heading("5. UX Blockers", level=1)
    table(
        doc,
        ["ID", "Severity", "Issue"],
        [
            ["B1", "CRITICAL", "iOS manual dive edit unreachable"],
            ["B2", "CRITICAL", "iCloud merge drops manual dive metadata"],
            ["B3", "CRITICAL", "Planner hardcoded ascent row misleads user"],
            ["H1", "HIGH", "Companion disclaimer every launch"],
            ["H5", "HIGH", "Watch runtime alarm 30 vs 60 default mismatch"],
            ["H6", "HIGH", "Watch units not applied on Live"],
        ],
    )

    doc.add_heading("6. Safety Issues", level=1)
    bullets(
        doc,
        [
            "Mock planner row undermines indicative-planning disclaimer (HIGH).",
            "Manual metadata loss on iCloud merge (HIGH).",
            "Ascent inline banner + depth limits + legal gates: strong (LOW risk).",
        ],
    )

    doc.add_heading("7. Recommended Priority", level=1)
    bullets(
        doc,
        [
            "Immediate: B1, B2, B3, H5, H2 (manual pressures UI).",
            "Pre-release: H1, H3 share, H4 pairing reset, H6 units display.",
            "Post-release: App Intents catalog, dead GPS views, i18n cleanup.",
        ],
    )

    doc.add_heading("8. Code Impact", level=1)
    para(
        doc,
        "Most fixes are small UI/wiring (navigation link, merge fields, delete mock row, defaults). "
        "Medium: unit formatting across Watch views; planner export share. No architectural rewrite required.",
    )

    doc.add_heading("Reference UI", level=1)
    if IMG_WATCH.exists():
        doc.add_picture(str(IMG_WATCH), width=Inches(2.2))
        para(doc, "Watch reference", size=9)
    if IMG_IOS.exists():
        doc.add_picture(str(IMG_IOS), width=Inches(2.2))
        para(doc, "iOS reference", size=9)

    para(doc, f"Full markdown source: {MD.name}", size=9)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
