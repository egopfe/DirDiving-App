#!/usr/bin/env python3
"""Render the 2026-05-19 MAIN UX audit Markdown into a clean .docx.

Usage:
    python3 Docs/generate_main_ux_audit_20260519_docx.py

Reads:
    Docs/MAIN_BRANCHES_UX_INTERACTION_AUDIT_20260519_CURRENT_PRE_MODIFICATION.md

Writes:
    Docs/MAIN_BRANCHES_UX_INTERACTION_AUDIT_20260519_CURRENT_PRE_MODIFICATION.docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


HERE = Path(__file__).resolve().parent
SRC = HERE / "MAIN_BRANCHES_UX_INTERACTION_AUDIT_20260519_CURRENT_PRE_MODIFICATION.md"
OUT = HERE / "MAIN_BRANCHES_UX_INTERACTION_AUDIT_20260519_CURRENT_PRE_MODIFICATION.docx"


def add_inline_runs(paragraph, text: str) -> None:
    """Render very small subset of markdown inline syntax (bold, code, italic)."""
    # tokens are split by **bold**, `code`, *italic*
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(9.5)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def render(md_text: str) -> Document:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        raw = lines[i].rstrip()

        # markdown table
        if raw.startswith("|") and i + 1 < len(lines) and re.match(r"\|[\s:|-]+\|", lines[i + 1]):
            header = [cell.strip() for cell in raw.strip("|").split("|")]
            i += 2  # skip header separator
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                if len(cells) < len(header):
                    cells.extend([""] * (len(header) - len(cells)))
                elif len(cells) > len(header):
                    cells = cells[: len(header) - 1] + ["|".join(cells[len(header) - 1 :])]
                rows.append(cells)
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            table.style = "Light Grid Accent 1"
            for col_idx, value in enumerate(header):
                cell = table.rows[0].cells[col_idx]
                cell.text = ""
                add_inline_runs(cell.paragraphs[0], value)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row_idx, row in enumerate(rows, start=1):
                for col_idx, value in enumerate(row):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = ""
                    add_inline_runs(cell.paragraphs[0], value)
            doc.add_paragraph("")
            continue

        if raw.startswith("# "):
            heading = doc.add_heading(level=0)
            add_inline_runs(heading, raw[2:].strip())
        elif raw.startswith("## "):
            heading = doc.add_heading(level=1)
            add_inline_runs(heading, raw[3:].strip())
        elif raw.startswith("### "):
            heading = doc.add_heading(level=2)
            add_inline_runs(heading, raw[4:].strip())
        elif raw.startswith("#### "):
            heading = doc.add_heading(level=3)
            add_inline_runs(heading, raw[5:].strip())
        elif raw.startswith("- ") or raw.startswith("* "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, raw[2:].strip())
        elif re.match(r"^\d+\.\s", raw):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_runs(paragraph, re.sub(r"^\d+\.\s", "", raw, count=1))
        elif raw.startswith("---"):
            doc.add_paragraph().add_run().add_break()
        elif raw.strip() == "":
            doc.add_paragraph("")
        else:
            paragraph = doc.add_paragraph()
            add_inline_runs(paragraph, raw)

        i += 1

    return doc


def main() -> None:
    md_text = SRC.read_text(encoding="utf-8")
    doc = render(md_text)
    doc.save(OUT)
    print(f"Wrote {OUT.relative_to(HERE.parent)}")


if __name__ == "__main__":
    main()
