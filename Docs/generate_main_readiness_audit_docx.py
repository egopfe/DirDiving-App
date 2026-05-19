"""Generate MAIN branch readiness audit Word document. Run: python Docs/generate_main_readiness_audit_docx.py"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent / "MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260519.docx"


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def main():
    doc = Document()
    title = doc.add_heading("DIR DIVING — MAIN BRANCH COMPLETE READINESS AUDIT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(
        doc,
        "Audit scope: branch `main` only — Apple Watch target `DIRDiving Watch App` + iOS target `DIRDiving iOS` "
        "as defined in project.yml. Experimental branches were not reviewed. No code was modified for this audit. "
        "Build commands (xcodegen / xcodebuild) were not available in this environment (Windows; no Xcode in PATH).",
    )
    add_para(doc, "Reference benchmarks: premium black/neon Watch UI; dark marine cyan-accent iOS companion (per product documentation and screenshots referenced in README/Docs).")
    add_para(doc, "Lingua: questo report e redatto principalmente in inglese per allineamento al comando di audit; le etichette UI citate restano quelle del codice (italiano/inglese misto).")

    # Detailed sections 1-14 (audit checklist)
    add_heading(doc, "Checklist dettagliata (sezioni 1-14)", 1)

    add_heading(doc, "1. Branch and build verification", 2)
    add_bullets(
        doc,
        [
            "Branch: `main` confermato.",
            "Watch: nessun import runtime verso file esclusi da project.yml per il target Watch (Apnea/Snorkeling/Buddy esclusi).",
            "iOS: dipendenze rotte — vedi A e M.",
            "project.yml: struttura YAML coerente (nomi target, bundle ID com.egopfe.dirdiving / .ios, entitlements paths, embed Watch in iOS).",
            "XcodeGen: non eseguito (tool assente su host).",
            "xcodebuild: non eseguito.",
            "Bundle ID: WK companion com.egopfe.dirdiving.ios allineato a iOS bundle.",
            "Entitlements: Watch include water-submersion + iCloud KVS/CloudKit container; iOS include iCloud senza water-submersion (atteso).",
            "WatchConnectivity: framework dichiarato; pairing dipende da app iPhone.",
            "Asset iOS AppIcon: Contents.json referenzia PNG presenti nella cartella appiconset (verifica statica OK).",
            "TODO bloccanti: mismatch iOS compile e priorita superiore a TODO di copy.",
        ],
    )

    add_heading(doc, "2. Apple Watch MAIN feature audit", 2)
    add_para(doc, "Copertura: schermata live, profondita, runtime, TTV, cronometro, profondita media/max, temperatura, gauge risalita, avviso risalita, bussola, log/dettaglio, GPS inizio/fine, export, immagini, impostazioni, allarmi, info/batteria, unita, aptica. Stato sintetico: implementato e raggiungibile per il core Diving; validazione sensore/entitlement su hardware.")

    add_heading(doc, "3. iOS Companion MAIN feature audit", 2)
    add_para(doc, "Il codice sorgente iOS contiene Logbook, DiveDetail, Planner, Analysis, Equipment, More, sync — ma il target `DIRDiving iOS` su `main` non e attualmente compilabile; quindi reachability/usability non verificabili in runtime su questo branch.")

    add_heading(doc, "4. UI consistency audit", 2)
    add_bullets(
        doc,
        [
            "Watch: palette nera + cyan/green/yellow/red + pannelli arrotondati — coerente con linea guida README.",
            "Watch: verificare sovrapposizioni solo su device (non misurabile qui).",
            "iOS: tema DIRTheme coerente; card experimental in Analysis da rivalutare per policy 'MAIN senza placeholder'.",
        ],
    )

    add_heading(doc, "5. UX completeness audit", 2)
    add_bullets(
        doc,
        [
            "Watch: flusso pre-immersione -> live -> log ragionevole per utente medio; selector indica solo Diving stabile.",
            "Watch: export disabilitato/etichettato con log vuoto (empty state).",
            "iOS: bloccato da compile — workflow utente non testabile.",
        ],
    )

    add_heading(doc, "6. Settings audit", 2)
    add_para(doc, "Watch: copertura buona (risalita, allarmi, aptica, sync, GPS, export note). iOS: da validare dopo fix compile (MoreView).")

    add_heading(doc, "7. Haptics and tones audit", 2)
    add_para(doc, "Watch: aptica presente per allarmi e risalita e cronometro; toni audio custom non individuati. iOS: notifiche/permessi documentati in README; nessun sistema toni in-app auditato.")

    add_heading(doc, "8. Apple Watch hardware interaction audit", 2)
    add_para(doc, "TabView verticale + Crown; long press/confirmation dialog su delete; Action Button solo via intents — nessuna promessa di callback hardware dedicato nel flusso principale.")

    add_heading(doc, "9. Sync audit", 2)
    add_para(doc, "Watch->iPhone: code paths per messaggio diretto e coda; stati pending/sent/ack/failed esposti; ack richiesto prima di rimuovere pending (per design documentato). iPhone->Watch settings: non production-complete.")

    add_heading(doc, "10. Export audit", 2)
    add_para(doc, "CSV Subsurface-centric; Watch share sheet pattern; iOS export service presente nel tree. GPX/KML non come deliverable MAIN obbligatorio.")

    add_heading(doc, "11. Safety / disclaimer audit", 2)
    add_para(doc, "Non dichiarato computer subacqueo certificato in documentazione; TTV chiarito; profondita richiede validazione Apple; planner iOS disclaimer da verificare quando buildabile.")

    add_heading(doc, "12. Error and empty state audit", 2)
    add_para(doc, "Watch: stati vuoti e messaggi GPS/sync; DiveManager gestisce indisponibilita sensore. iOS: non eseguito.")

    add_heading(doc, "13. Code quality / self-consistency audit", 2)
    add_bullets(
        doc,
        [
            "Inconsistenza critica: esclusioni XcodeGen vs entry point iOS.",
            "HapticService include API Buddy non usate in UI MAIN — piccolo debito tecnico.",
            "Nessuna importazione accidentalmente da ExplorationStore nel target Watch (escluso).",
        ],
    )

    add_heading(doc, "14. Report format", 2)
    add_para(doc, "Sezioni A-O seguono il formato richiesto; tabella feature in C.")

    # A
    add_heading(doc, "A. Branch Confirmed", 1)
    add_bullets(
        doc,
        [
            "Current branch audited: `main` (local; ahead of origin by prior commits).",
            "Targets: `DIRDiving Watch App` (watchOS 10.0), `DIRDiving iOS` (iOS 17.0), schemes per project.yml.",
            "Build status: NOT VERIFIED HERE — xcodegen / xcodebuild absent on host. Static analysis shows CRITICAL iOS compile mismatch (see B, M).",
            "Watch: project.yml excludes experimental-only sources (Apnea, Snorkeling, BuddyAssist, ExplorationStore, etc.); Watch app does not require those files.",
            "iOS: project.yml excludes `PlannerStore.swift`, `ExplorationPlanningStore.swift`, `BuddyExperimentalStore.swift`, `ExplorationCenterView.swift`, `BuddyExperimentalView.swift`, "
            "but `DIRDivingiOSApp.swift` and `ContentView.swift` still reference those types — target will not compile until project.yml or app entry is reconciled.",
        ],
    )

    # B
    add_heading(doc, "B. Executive Summary", 1)
    add_para(doc, "Percentages are engineering estimates from static review (not device-tested).", bold=True)
    add_bullets(
        doc,
        [
            "Overall readiness (dual-target product as on `main`): ~45% — blocked by iOS target compile failure.",
            "Apple Watch MAIN readiness: ~82% — structure and UX look coherent; needs macOS/Xcode build + Apple Watch Ultra validation.",
            "iOS Companion MAIN (in-repo): ~30% — compile broken on `main` as configured; logic may be fine once sources/targets align.",
            "UX readiness (Watch): ~78% — flows mostly clear; mode selector is Diving-only (good); minor experimental-flavored copy on iOS Analysis if that file ships.",
            "Safety / disclaimer readiness: ~76% — TTV documented as non-deco in Watch UI/accessibility; README covers non-certification; depth entitlement still needs Apple validation.",
            "Compile readiness (full XcodeGen project): FAILED for iOS target as currently wired.",
        ],
    )

    # C - table
    add_heading(doc, "C. Feature Inventory (MAIN)", 1)
    add_para(doc, "Legend: Y=yes/present, P=partial, N=no, ?=cannot verify without build/device.")

    rows_watch = [
        ("Watch", "Live dive screen", "Y", "Y", "Y", "Y", "Premium panels; haptics-off banner when disabled", "Low"),
        ("Watch", "Depth display", "Y", "Y", "Y", "P", "Depends on CMWaterSubmersion + entitlement; simulator/device limited", "Med"),
        ("Watch", "Runtime", "Y", "Y", "Y", "Y", "", "Low"),
        ("Watch", "TTV", "Y", "Y", "Y", "Y", "Informative metric (avg depth + runtime); not NDL/TTS; labeled in settings/a11y", "Low"),
        ("Watch", "Stopwatch start/stop/reset", "Y", "Y", "Y", "Y", "Haptics on start/stop/reset when enabled", "Low"),
        ("Watch", "Avg / max depth", "Y", "Y", "Y", "Y", "", "Low"),
        ("Watch", "Temperature", "Y", "Y", "P", "P", "When sensor nil, may show no reading — not same as fake zero in newer iOS builds", "Low"),
        ("Watch", "Ascent-rate gauge", "Y", "Y", "Y", "Y", "", "Low"),
        ("Watch", "Ascent warning overlay", "Y", "Y", "Y", "Y", "Context + Formatters.zero path present in codebase audited", "Low"),
        ("Watch", "Bussola / bearing", "Y", "Y", "Y", "Y", "SET BEARING / CLEAR; location permission", "Low"),
        ("Watch", "Dive log list", "Y", "Y", "Y", "Y", "Empty state + delete confirmation", "Low"),
        ("Watch", "Dive detail", "Y", "Y", "Y", "Y", "Export navigation", "Low"),
        ("Watch", "GPS start / end", "Y", "Y", "Y", "Y", "Surface-only; fix vs fallback vs no-fix modeled", "Med"),
        ("Watch", "CSV export / share", "Y", "Y", "Y", "Y", "Subsurface-oriented; metric", "Low"),
        ("Watch", "User images", "Y", "Y", "Y", "Y", "Empty state when no bundled assets", "Low"),
        ("Watch", "Settings", "Y", "Y", "Y", "Y", "Ascent limits, alarms, haptics, sync queue, TTV note, imperial disabled", "Low"),
        ("Watch", "Units (imperial)", "P", "Y", "Y", "P", "Imperial not implemented — selection disabled / forced metric", "Low"),
        ("Watch", "Haptics", "Y", "Y", "Y", "Y", "Gated by AppStorage; buddy-specific haptics compiled but Buddy not in MAIN", "Low"),
        ("Watch", "Tones / sounds", "N", "-", "-", "-", "No custom alert tones; uses WK haptic patterns only", "Low"),
        ("Watch", "Digital Crown", "Y", "Y", "Y", "Y", "Vertical page TabView — standard Crown navigation", "Low"),
        ("Watch", "Side / Action Button", "P", "?", "?", "P", "App Intents exist in repo; no guaranteed hardware callback — docs should set expectations", "Med"),
    ]

    rows_ios = [
        ("iOS", "Logbook", "Y", "?", "?", "?", "Cannot compile target on main — inferred from sources", "CRITICAL"),
        ("iOS", "Dive detail / charts", "Y", "?", "?", "?", "Same — files present", "CRITICAL"),
        ("iOS", "Planner / Plan result", "Y", "?", "?", "?", "PlannerStore excluded but PlannerView requires it", "CRITICAL"),
        ("iOS", "Gas / Bühlmann UI", "Y", "?", "?", "?", "Tied to PlannerStore", "CRITICAL"),
        ("iOS", "Export / Subsurface", "Y", "?", "?", "?", "SubsurfaceExportService in tree", "CRITICAL"),
        ("iOS", "Watch sync", "Y", "?", "?", "?", "WatchSyncService in tree; ack path documented in prior fixes", "CRITICAL"),
        ("iOS", "Analysis", "Y", "?", "?", "P", "Includes 'EXPERIMENTAL ANALYTICS CONCEPTS' mock card — may violate 'no placeholder in MAIN' product goal", "Med"),
        ("iOS", "Settings / More", "Y", "?", "?", "?", "MoreView in tree", "CRITICAL"),
        ("iOS", "Explore / Buddy tabs", "Y", "N", "N", "N", "In ContentView but sources excluded — compile error", "CRITICAL"),
    ]

    table = doc.add_table(rows=1, cols=8)
    hdr = table.rows[0].cells
    headers = ["Platform", "Feature", "Implemented", "Reachable", "Usable", "Complete", "Notes", "Severity"]
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows_watch + rows_ios:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    add_heading(doc, "D. Navigation Map", 1)
    add_para(doc, "Watch MAIN (TabView pages): Mode selector → Live → Compass → Settings → User images → Dive log. Vertical paging = Crown-friendly.")
    add_para(doc, "iOS MAIN (intended): Tab bar; current ContentView references excluded tabs — navigation inconsistent with build graph until fixed.")
    add_bullets(doc, ["Dead end risk: low on Watch for core dive loop.", "Dead end risk: HIGH on iOS — app will not launch from this branch state."])

    add_heading(doc, "E. UI Consistency vs Reference", 1)
    add_bullets(
        doc,
        [
            "Watch: DiveUIComponents / DiveScreenBackground / neon panels align with documented premium Watch reference.",
            "Watch: AscentGaugeView.swift shows local modification (line endings / uncommitted) — verify no visual regression in Xcode preview.",
            "iOS: DIRTheme / DIRBackground pattern matches dark cyan companion style when buildable.",
            "iOS: Analysis screen includes explicit mock/experimental analytics tiles — severity MEDIUM for 'average user' confusion if shipped.",
        ],
    )

    add_heading(doc, "F. Settings Report", 1)
    add_bullets(
        doc,
        [
            "Watch: ascent limits, alarms, haptics toggle, GPS/sync status rows, TTV disclaimer, export note, sync queue controls — reachable via Settings tab.",
            "Watch: units — metric enforced; imperial disabled (honest).",
            "Watch: settings sync — labeled local-only / planned (no false promise).",
            "iOS: cannot execute runtime audit; code suggests MoreView for permissions, cloud, Watch trust, notifications.",
        ],
    )

    add_heading(doc, "G. Haptics / Tones", 1)
    add_bullets(
        doc,
        [
            "Watch: dive alarms use warnIfNeeded (throttled); ascent over-limit uses same; stopwatch uses confirm/notify.",
            "Watch: no dive-start confirm haptic in beginDiveIfNeeded snippet — optional product gap (LOW).",
            "Watch: no dedicated haptic on GPS save (GPS is lifecycle-driven — LOW).",
            "iOS: notification permission flow in Settings (per docs); no custom in-app tones audited here.",
        ],
    )

    add_heading(doc, "H. Hardware Controls", 1)
    add_bullets(
        doc,
        [
            "Crown: vertical TabView — supported.",
            "Long press: used for dive delete context where implemented; list also has explicit delete + confirmation.",
            "Action Button: rely on App Intents / system exposure — document limitation (MEDIUM clarity).",
        ],
    )

    add_heading(doc, "I. Sync Report", 1)
    add_bullets(
        doc,
        [
            "Watch → iPhone: sendMessage + transferUserInfo; pending/sent/acknowledged/failed counts in settings (truthful queue UX).",
            "iPhone → Watch: not a full settings sync pipeline; documented as local/planned.",
            "Duplicate prevention: session id / codec (per existing services — device QA required).",
        ],
    )

    add_heading(doc, "J. Export Report", 1)
    add_bullets(
        doc,
        [
            "Watch: CSV export path from log + ExportView; empty log disables export messaging.",
            "iOS: Subsurface CSV + import service in tree; fix-source columns per docs.",
            "GPX/KML: not primary in MAIN scope.",
        ],
    )

    add_heading(doc, "K. Safety / Disclaimer", 1)
    add_bullets(
        doc,
        [
            "App positioning: documentation states not a certified dive computer — keep visible in-app on all planner/dive surfaces (verify on iOS when buildable).",
            "TTV: non-decompression messaging present on Watch live/settings/a11y.",
            "Depth: entitlement present in plist; Apple portal + hardware validation still required.",
            "App Store: privacy strings present for motion + location on Watch Info.plist.",
        ],
    )

    add_heading(doc, "L. Error / Empty States", 1)
    add_bullets(
        doc,
        [
            "Watch: empty logbook, UserImages empty state, GPS no-fix paths, sync failure counts.",
            "Simulator: depth may be unavailable — error string in DiveManager.",
        ],
    )

    add_heading(doc, "M. Bugs To Fix", 1)
    bugs = [
        {
            "title": "iOS target compile failure (excluded sources still referenced)",
            "platform": "iOS",
            "where": "project.yml (excludes); iOSApp/App/DIRDivingiOSApp.swift; iOSApp/Views/ContentView.swift",
            "severity": "CRITICAL",
            "impact": "XcodeGen project cannot build `DIRDiving iOS`; no TestFlight/App Store; companion unusable from this branch.",
            "fix": "Either remove tabs and @StateObject for excluded stores/views, OR remove those paths from project.yml excludes so sources compile; align with intended MAIN tab set (Logbook/Explore/Analysis/Planner/Gear/Settings).",
            "effort": "small functional to medium (yml + app entry + ContentView only if stores stay).",
        },
        {
            "title": "MAIN tab bar exposes experimental labels without matching MAIN product policy",
            "platform": "iOS",
            "where": "iOSApp/Views/ContentView.swift",
            "severity": "HIGH",
            "impact": "Average user sees Buddy Lab / Explore Lab on a production-oriented branch definition.",
            "fix": "Match `main-iOS` style tab set or clearly gate labs behind feature flags removed from App Store builds.",
            "effort": "UI-only + small navigation.",
        },
        {
            "title": "Experimental analytics mock card in Analysis",
            "platform": "iOS",
            "where": "iOSApp/Views/AnalysisView.swift (adaptiveAnalyticsConcept)",
            "severity": "MEDIUM",
            "impact": "Looks like unfinished or misleading health/readiness metrics.",
            "fix": "Remove card for MAIN or replace with real metrics only; add explicit LAB label if retained in internal builds.",
            "effort": "UI-only.",
        },
        {
            "title": "No haptic on automatic dive start",
            "platform": "Watch",
            "where": "Services/DiveManager.swift (beginDiveIfNeeded)",
            "severity": "LOW",
            "impact": "User may not notice session start without visual attention.",
            "fix": "Optional notify() when depth automation starts session.",
            "effort": "small functional.",
        },
        {
            "title": "Buddy haptic helpers compiled without Buddy MAIN UI",
            "platform": "Watch",
            "where": "Services/HapticService.swift",
            "severity": "LOW",
            "impact": "Dead code; maintenance noise.",
            "fix": "Wrap in #if or move to experimental target only.",
            "effort": "small refactor / conditional compile.",
        },
        {
            "title": "Dirty working tree file AscentGaugeView.swift",
            "platform": "Watch",
            "where": "Views/AscentGaugeView.swift",
            "severity": "LOW",
            "impact": "CI noise or accidental wrong commit.",
            "fix": "Normalize line endings or revert if no intentional change.",
            "effort": "UI-only / housekeeping.",
        },
    ]
    for b in bugs:
        add_para(doc, f"{b['title']} — {b['severity']} ({b['platform']})", bold=True)
        add_para(doc, f"File/screen: {b['where']}")
        add_para(doc, f"User impact: {b['impact']}")
        add_para(doc, f"Recommended fix: {b['fix']}")
        add_para(doc, f"Estimated code impact: {b['effort']}")
        doc.add_paragraph()

    add_heading(doc, "N. Priority Roadmap", 1)
    add_para(doc, "1. Must fix before compile/use: Align iOS target sources with DIRDivingiOSApp + ContentView (include stores/views OR remove tabs and @StateObject deps).")
    add_para(doc, "2. Must fix before TestFlight: Full macOS xcodegen + archive; Watch Ultra depth + GPS QA; iOS Watch sync QA; remove or gate mock analytics card.")
    add_para(doc, "3. Must fix before App Store: App Store metadata, support URL, encryption, screenshots, entitlement proofs, crash-free session.")
    add_para(doc, "4. Post-release: Action Button discoverability, optional tones, richer conflict UI.")

    add_heading(doc, "O. Final Verdict", 1)
    add_bullets(
        doc,
        [
            "Ready to compile (full repo): NO — iOS target broken on `main` as configured.",
            "Ready for internal test (Watch-only local build after fixing iOS or excluding iOS target): PARTIAL — needs Xcode build.",
            "Ready for average user: NO — iOS companion not shippable from this branch state.",
            "Ready for TestFlight: NO.",
            "Ready for App Store: NO.",
            "What blocks 100%: iOS compile graph mismatch; lack of host-side Xcode build verification; hardware validation for depth/GPS/sync.",
        ],
    )

    add_heading(doc, "Appendice", 1)
    add_para(
        doc,
        "Rigenerazione documento: eseguire `python Docs/generate_main_readiness_audit_docx.py` dalla root del repository. "
        f"Output: `{OUT.name}`.",
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
