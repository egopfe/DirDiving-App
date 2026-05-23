# -*- coding: utf-8 -*-
"""Generate MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260523.docx from markdown summary."""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
OUT = HERE / "MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260523.docx"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"


def _git(*args: str) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=REPO, text=True).strip()
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        return "n/d"


def para(doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    doc = Document()
    head = _git("rev-parse", "HEAD")
    branch = _git("rev-parse", "--abbrev-ref", "HEAD")

    title = doc.add_heading("DIR DIVING — MAIN BRANCH COMPLETE READINESS AUDIT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, f"Date: 2026-05-23 · Branch: {branch} · HEAD: {head}", bold=True)
    para(
        doc,
        "Audit-only: no code changes. Scope: Apple Watch MAIN + iOS Companion MAIN on unified main. "
        "Visual benchmarks: Docs/ReferenceUI/*.png.",
    )

    doc.add_heading("A. Branch Confirmed", level=1)
    bullets(
        doc,
        [
            f"Branch: {branch} · HEAD: {head}",
            "project.yml valid; experimental sources excluded from MAIN targets.",
            "xcodegen generate: PASS.",
            "xcodebuild Watch (Ultra 3 sim), iOS (iPhone 17), generic iOS Simulator: BUILD SUCCEEDED.",
            "Bundle IDs: iOS com.egopfe.dirdiving.ios; Watch com.egopfe.dirdiving.ios.watch.",
            "Entitlements: Watch water submersion + iCloud KVS; iOS iCloud KVS only.",
        ],
    )

    doc.add_heading("B. Executive Summary", level=1)
    table(
        doc,
        ["Dimension", "Readiness %"],
        [
            ["Overall MAIN", "86%"],
            ["Apple Watch MAIN", "88%"],
            ["iOS Companion MAIN", "87%"],
            ["UX completeness", "82%"],
            ["Safety / disclaimers", "88%"],
            ["Compile readiness", "97%"],
        ],
    )
    para(
        doc,
        "Verdict: Ready to compile and for internal/paired-device QA. Not 100% for average consumer or App Store "
        "without Ultra depth entitlement on com.egopfe.dirdiving.ios.watch, physical sync QA, and store assets.",
    )

    doc.add_heading("C. Feature Inventory (summary)", level=1)
    para(doc, "Full table in MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260523.md.", size=9)
    table(
        doc,
        ["Area", "Status"],
        [
            ["Watch live dive + depth safety 35/38/40m", "Complete in code; hardware partial"],
            ["Watch compass, log, export, settings", "Complete / partial i18n"],
            ["Watch sync", "Implemented; device QA required"],
            ["iOS logbook, analysis, planner, equipment, more", "Mostly complete"],
            ["iOS sync push + conflict UI", "Implemented"],
            ["Experimental modes", "Excluded from MAIN binary"],
        ],
    )

    doc.add_heading("D–L. Detailed Sections", level=1)
    bullets(
        doc,
        [
            "D. Navigation: Watch vertical pages; iOS 5 tabs; no critical dead ends.",
            "E. UI: Matches reference screenshots; planner tab chrome partial on iOS.",
            "F. Settings: Local persistence; no Watch↔iOS settings sync.",
            "G. Haptics: Watch yes; audio tones intentionally absent.",
            "H. Hardware: Crown page navigation; Action Button partial.",
            "I. Sync: Bidirectional sessions; HMAC; conflict card on iOS.",
            "J. Export: Subsurface CSV Watch + iOS; import CSV iOS only.",
            "K. Safety: Legal gate, depth limits, planner ack; entitlement external.",
            "L. Empty/error states: Present on primary surfaces.",
        ],
    )

    doc.add_heading("M. Top Bugs To Fix", level=1)
    table(
        doc,
        ["Title", "Severity", "Impact"],
        [
            ["Ultra depth + entitlement not field-proven", "CRITICAL", "No real depth without Apple approval"],
            ["Physical sync/tombstone QA", "HIGH", "Trust in companion sync"],
            ["CSV import hidden when logbook non-empty", "MEDIUM", "Import discoverability"],
            ["Planner result tabs display-only", "MEDIUM", "Confusing UX"],
        ],
    )

    doc.add_heading("N. Priority Roadmap", level=1)
    bullets(
        doc,
        [
            "Before TestFlight: entitlement on .ios.watch, physical QA, legal/assets.",
            "Before App Store: all above + planner polish + marketing review.",
            "Post-release: settings cross-sync, Action Button catalog, dead code cleanup.",
        ],
    )

    doc.add_heading("O. Final Verdict", level=1)
    table(
        doc,
        ["Question", "Answer"],
        [
            ["Ready to compile?", "YES"],
            ["Ready for internal test?", "YES (with playbook)"],
            ["Ready for average user?", "CONDITIONAL"],
            ["Ready for TestFlight?", "CONDITIONAL"],
            ["Ready for App Store?", "NO"],
        ],
    )

    doc.add_heading("Reference UI", level=1)
    if IMG_WATCH.exists():
        para(doc, "Watch LIVE reference:", bold=True)
        doc.add_picture(str(IMG_WATCH), width=Inches(2.2))
    if IMG_IOS.exists():
        para(doc, "iOS Companion reference:", bold=True)
        doc.add_picture(str(IMG_IOS), width=Inches(2.8))

    para(doc, "Full markdown report: Docs/MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260523.md", size=9)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
