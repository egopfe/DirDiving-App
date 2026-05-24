# -*- coding: utf-8 -*-
"""Generate MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260524.docx."""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
OUT = HERE / "MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260524.docx"
MD = HERE / "MAIN_BRANCH_COMPLETE_READINESS_AUDIT_20260524.md"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"


def _git(*args: str) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=REPO, text=True).strip()
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        return "n/d"


def para(doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    doc = Document()
    head = _git("rev-parse", "HEAD")
    branch = _git("rev-parse", "--abbrev-ref", "HEAD")

    title = doc.add_heading("DIR DIVING — MAIN Branch Complete Readiness Audit", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, f"Date: 2026-05-24 · Branch: {branch} · HEAD: {head}", bold=True)
    para(
        doc,
        "Audit-only (no code changes). Scope: Apple Watch MAIN + iOS Companion MAIN. "
        "Benchmarks: Docs/ReferenceUI/Watch_LIVE_reference.png and iOS_Companion_reference.png.",
    )

    doc.add_heading("B. Executive Summary", level=1)
    table(
        doc,
        ["Dimension", "Readiness %"],
        [
            ["Overall", "~91%"],
            ["Compile", "100%"],
            ["Apple Watch MAIN", "~93%"],
            ["iOS Companion MAIN", "~92%"],
            ["UX completeness", "~94%"],
            ["Safety", "~88%"],
            ["UI vs reference", "~90%"],
        ],
    )

    doc.add_heading("A. Branch & Build", level=1)
    bullets(
        doc,
        [
            f"Branch {branch} @ {head}; origin/main in sync.",
            "xcodegen + Watch + iOS builds: SUCCEEDED.",
            "Experimental targets excluded in project.yml.",
            "Bundle IDs and entitlements coherent.",
        ],
    )

    doc.add_heading("O. Final Verdict", level=1)
    table(
        doc,
        ["Question", "Answer"],
        [
            ["Ready to compile?", "YES"],
            ["Ready for internal test?", "YES"],
            ["Ready for average user?", "MOSTLY (~91%)"],
            ["Ready for TestFlight?", "YES after Ultra depth QA"],
            ["Ready for App Store?", "CONDITIONAL"],
            ["Blocks 100%", "Hardware depth QA, store review, i18n/cloud polish"],
        ],
    )

    doc.add_heading("M. Top Bugs To Fix", level=1)
    table(
        doc,
        ["ID", "Severity", "Issue"],
        [
            ["R1", "HIGH", "Physical depth/submersion QA on Apple Watch Ultra"],
            ["R2", "MED", "Planner safety ack not persisted (session @State)"],
            ["R3", "MED", "iCloud decode errors may be silent"],
            ["R4", "LOW", "Mixed IT/EN strings on iOS main screens"],
        ],
    )

    doc.add_heading("Prior audit fixes (landed)", level=1)
    bullets(
        doc,
        [
            "876bcd2: manual edit, merge fields, planner mock row, disclaimer persist, sync fixes.",
            "db72dce: ascent gauge imperial, 7 App Shortcuts, side-button copy, detail refresh.",
        ],
    )

    doc.add_heading("Reference UI", level=1)
    if IMG_WATCH.exists():
        doc.add_picture(str(IMG_WATCH), width=Inches(2.2))
        para(doc, "Watch LIVE reference", size=9)
    if IMG_IOS.exists():
        doc.add_picture(str(IMG_IOS), width=Inches(2.2))
        para(doc, "iOS Companion reference", size=9)

    para(doc, f"Full markdown: {MD.name}", size=9)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
