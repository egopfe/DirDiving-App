# -*- coding: utf-8 -*-
"""Generate MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260523.docx."""
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = Path(__file__).resolve().parent
REPO = HERE.parent
OUT = HERE / "MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260523.docx"
IMG_WATCH = HERE / "ReferenceUI" / "Watch_LIVE_reference.png"
IMG_IOS = HERE / "ReferenceUI" / "iOS_Companion_reference.png"


def _git(*args: str) -> str:
    try:
        return subprocess.check_output(["git", *args], cwd=REPO, text=True).strip()
    except (subprocess.CalledProcessError, FileNotFoundError, OSError):
        return "n/d"


def para(doc: Document, text: str, bold: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def main() -> None:
    doc = Document()
    head = _git("rev-parse", "HEAD")
    branch = _git("rev-parse", "--abbrev-ref", "HEAD")

    title = doc.add_heading(
        "DIR DIVING — MAIN BRANCH UX / INTERACTION / FEATURE ACCESSIBILITY AUDIT",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, f"Date: 2026-05-23 · Branch: {branch} · HEAD: {head}", bold=True)
    para(
        doc,
        "Pre-modification audit only — no code changes. Scope: Apple Watch MAIN + iOS Companion MAIN. "
        "Experimental branches/files excluded.",
    )

    doc.add_heading("1. Feature Inventory (summary)", level=1)
    table(
        doc,
        ["Feature area", "Watch", "iOS"],
        [
            ["Live dive + depth safety 35/38/40 m", "Reachable, complete", "Review in logbook/detail"],
            ["Compass / BUSSOLA", "Reachable (Crown page)", "N/A on Watch-only surface"],
            ["Dive log + export", "Complete", "Complete + Subsurface export"],
            ["CSV import", "N/A", "Hidden when Analysis non-empty"],
            ["Planner", "N/A", "Calculate OK; tabs/mode misleading"],
            ["Settings", "Rich tab; export row non-actionable", "More tab; no units picker"],
            ["Sync Watch↔iPhone", "Status + retry", "Push + conflict UI"],
            ["Mode Selection", "Implemented, hidden", "N/A"],
            ["HR / notifications", "Not in MAIN", "Not in MAIN"],
        ],
    )

    doc.add_heading("2. Navigation Map", level=1)
    para(doc, "Watch (Digital Crown vertical pages):", bold=True)
    bullets(
        doc,
        [
            "Legal gate → LIVE ↔ COMPASS ↔ SETTINGS ↔ [USER IMAGES if assets] ↔ DIVE LOG",
            "Settings children: Ascent, Alarms, Legal, Shortcuts help, Info — all return via back",
            "DIVE LOG → Detail → ExportView; delete with confirmation",
            "Mode Selection: unreachable (hasMultipleStableModes = false)",
        ],
    )
    para(doc, "iOS (5 tabs):", bold=True)
    bullets(
        doc,
        [
            "Legal gate → Logbook | Analysis | Planner | Equipment | More",
            "Logbook → DiveDetail; Planner → PlanResult on calculate",
            "No deep links; CSV import only from empty Analysis",
        ],
    )

    doc.add_heading("3. Settings Report", level=1)
    table(
        doc,
        ["Gap", "Platform", "Severity"],
        [
            ["Export row looks tappable but is informational", "Watch", "MEDIUM"],
            ["Units @AppStorage with no UI picker", "iOS", "MEDIUM"],
            ["Alarms / haptics only on Watch", "Both", "LOW (by design)"],
            ["No Watch↔iOS settings sync", "Both", "LOW post-release"],
            ["Depth safety thresholds not user-configurable", "Watch", "LOW (fixed policy)"],
        ],
    )

    doc.add_heading("4. Hardware Interaction Report", level=1)
    bullets(
        doc,
        [
            "Digital Crown: vertical page navigation + scroll in lists — OK",
            "Side button / Action Button: only stopwatch intents in catalog; dive/bearing need on-screen or user Shortcuts",
            "Long press: not used on MAIN (context menu delete on log only)",
            "Haptics: ascent, depth 35/38/40, alarms, GPS, export, compass set — gated by master toggle",
            "iOS: no haptic layer for sync/export",
        ],
    )

    doc.add_heading("5. UX Blockers", level=1)
    table(
        doc,
        ["Blocker", "Severity"],
        [
            ["CSV import only when Analysis empty", "HIGH"],
            ["Simulator depth requires manual dive (user confusion)", "HIGH"],
            ["Planner result tabs display-only", "MEDIUM"],
            ["Planner mode picker does not affect calculation", "MEDIUM"],
            ["Watch Settings Export row non-actionable", "MEDIUM"],
            ["iOS units preference has no UI", "MEDIUM"],
            ["Legal scroll-to-bottom honor system", "MEDIUM"],
            ["Decorative logbook +/⋯ buttons", "LOW"],
            ["Plan share button inert", "LOW"],
            ["resetPairingTrust not in UI", "LOW"],
        ],
    )

    doc.add_heading("6. Safety Issues (UX)", level=1)
    bullets(
        doc,
        [
            "Depth limit 35/38/40 m UI + throttled haptics: mitigated",
            "Ascent alarm + acknowledge banner: OK",
            "Planner requires safety acknowledgment: OK",
            "Silent sync when WC inactive: MEDIUM — status exists but easy to miss",
            "User depth alarm at 40 m may overlap safety exceeded state: LOW",
        ],
    )

    doc.add_heading("7. Recommended Priority Order", level=1)
    para(doc, "Immediate:", bold=True)
    bullets(
        doc,
        [
            "CSV import entry in More or Logbook",
            "Fix or label planner tabs and mode picker",
            "Watch Settings Export row — link or remove chevron",
            "iOS units picker in More",
        ],
    )
    para(doc, "Pre-release (TestFlight):", bold=True)
    bullets(
        doc,
        [
            "Legal scroll detection (optional)",
            "Expose Watch pairing reset in More (advanced)",
            "Physical device QA for sync and depth",
        ],
    )
    para(doc, "Post-release:", bold=True)
    bullets(
        doc,
        [
            "Expand App Intents catalog",
            "Settings cross-sync if product requires",
            "Mode Selection product decision",
        ],
    )

    doc.add_heading("8. Code Impact Report", level=1)
    table(
        doc,
        ["Fix", "Estimate"],
        [
            ["Import button, units picker, export row", "Small UI / small functional"],
            ["Planner tab wiring or labels", "UI-only or small"],
            ["Settings cross-sync", "Medium refactor"],
            ["No architectural UX blocker on MAIN", "—"],
        ],
    )

    doc.add_heading("9. Final Summary", level=1)
    table(
        doc,
        ["Metric", "Estimate"],
        [
            ["Release readiness (UX)", "82%"],
            ["UX completeness (Watch)", "84%"],
            ["UX completeness (iOS)", "80%"],
            ["Interaction completeness", "78%"],
            ["Stability (UX flows)", "85%"],
            ["Safety UX completeness", "86%"],
        ],
    )
    para(
        doc,
        "Verdict: Core dive → log → export on Watch and sync/review on iPhone are reachable for informed users. "
        "Not at App Store UX bar until CSV import discoverability and misleading planner chrome are addressed. "
        "Audit-only — no repository code modified.",
    )

    doc.add_heading("Reference UI", level=1)
    if IMG_WATCH.exists():
        para(doc, "Watch LIVE reference:", bold=True)
        doc.add_picture(str(IMG_WATCH), width=Inches(2.2))
    if IMG_IOS.exists():
        para(doc, "iOS Companion reference:", bold=True)
        doc.add_picture(str(IMG_IOS), width=Inches(2.8))

    para(
        doc,
        "Full markdown: Docs/MAIN_BRANCH_UX_INTERACTION_ACCESSIBILITY_AUDIT_20260523.md",
        size=9,
    )
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
